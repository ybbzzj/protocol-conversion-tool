from docx import Document
import os

def analyze_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File {file_path} not found.")
        return

    doc = Document(file_path)
    print(f"Analyzing {file_path}")
    print(f"Number of tables: {len(doc.tables)}")

    for i, table in enumerate(doc.tables):
        print(f"\nTable {i}:")
        print(f"Rows: {len(table.rows)}, Cols: {len(table.columns)}")
        
        # Print first few rows to see header
        for r_idx, row in enumerate(table.rows[:5]):
            cells = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
            # Handle duplicate cells caused by merged columns in python-docx
            unique_cells = []
            if cells:
                unique_cells.append(cells[0])
                for j in range(1, len(cells)):
                    if cells[j] != cells[j-1]:
                        unique_cells.append(cells[j])
            print(f"  Row {r_idx}: {unique_cells}")

analyze_docx("协议模板（公开）.docx")
analyze_docx("word/协议模板（公开）.docx")
