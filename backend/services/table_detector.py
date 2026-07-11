# -*- coding: utf-8 -*-
import logging
import re
from typing import List, Dict, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

# ─── 常量定义 ────────────────────────────────────────────────────────────────

# 信息名称行中不是消息名称的固定噪声词
INFO_NAME_ROW_NOISE = {
    '信息名称', '通信帧名称', '信息标识', '上级信息名称',
    '信息流向', '—', '－', '-', '', 'xx', 'XX',
    '信源、信宿', '信源、信目', '信源', '信宿',
    '传输周期', '发起时机', '错误处理', '其他',
}

# 数据类型关键字（用于识别字段定义表）
DATA_TYPE_KEYWORDS = {
    'UINTEGER', 'UINT', 'USHORT', 'UFLOAT', 'FLOAT',
    'DOUBLE', 'CHAR', 'BYTE', 'SHORT', 'BIT',
    'UINT8', 'UINT16', 'UINT32', 'INT8', 'INT16', 'INT32',
    '字节',
}

# 干扰表格前置段落关键词
NOISE_PARA_MARKERS = ['干扰表格', '测试用', '不需要提取']

# 干扰表格表头关键词
NOISE_HEADER_KEYWORDS = ['测试指令']

# 示例/目标格式表标志（首行包含这些列名，说明是示例表，跳过）
EXAMPLE_TABLE_HEADERS = {'名称', '信源系统码', '内容', '转换类型', '判读公式'}

# 内容列候选名（按优先级排列）
COL_CONTENT_CANDIDATES = ['内容', '数据含义', '字段', '参数', '信号名称', '名称']
# 数据类型列候选名
COL_TYPE_CANDIDATES = ['数据类型', '类型', '数据格式']
# 字节数列候选名
COL_BYTES_CANDIDATES = ['字节数', '字节', '数据长度（字节）', '长度']
# 值域列候选名
COL_RANGE_CANDIDATES = ['值域', '取值范围', '区间']
# 转换公式列候选名
COL_FORMULA_CANDIDATES = ['数据处理', '数据转换方法', '转换公式', '数据处理方法']
# 备注列候选名
COL_REMARK_CANDIDATES = ['备注', '说明', '数据来源']
# 单位列候选名
COL_UNIT_CANDIDATES = ['单位']
# 固定值列候选名（计算结果表中的"值"列）
COL_VALUE_CANDIDATES = ['值']

# Type A 候选行优先扫描词（含这些词的行优先作为候选行，但不再强制要求）
_TYPE_A_PREFERRED_SCAN_KEYWORDS = {
    '序号', '内容', '参数', '长度', '单位', '说明',
    '数据类型', '类型', '值域', '字节', '字节数'
}

# Type A 元数据行特征词（用于识别并跳过元数据行，避免误把元数据行当字段表头）
_TYPE_A_META_KEYWORDS = {
    '信息名称', '通信帧名称', '信息标识', '上级信息名称',
    '信息流向', '传输周期', '发起时机', '错误处理',
    '传送周期', '发送周期', '组成分组', '其他'
}


# ─── 辅助函数 ─────────────────────────────────────────────────────────────────

def _get_para_text(para_elem) -> str:
    """从段落 XML 元素中提取纯文本"""
    text = ''
    for child in para_elem:
        ctag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if ctag == 'r':
            for t in child:
                ttag = t.tag.split('}')[-1] if '}' in t.tag else t.tag
                if ttag == 't' and t.text:
                    text += t.text
    return text.strip()


def _build_grid(table) -> Tuple[List[List[str]], List[List[bool]]]:
    """
    用 python-docx 构建表格文本网格，精确处理合并单元格。

    - 水平合并（gridSpan）：同行跨列的单元格，grid 中每列都填写相同内容
    - 垂直合并（vMerge）：跨行合并，续行单元格填写起始行的内容，并标记 is_vmerge_cont=True

    返回：
        grid[r][c]         : 字符串，单元格文本（合并区域均填充实际值）
        is_vmerge_cont[r][c]: bool，True 表示该格是垂直合并的续行（非起始行）
    """
    n_rows = len(table.rows)
    # 列数用第一行的列数估算（含合并）
    try:
        n_cols = len(table.columns)
    except Exception:
        n_cols = max(len(r.cells) for r in table.rows) if n_rows else 0

    grid = [[''] * n_cols for _ in range(n_rows)]
    is_vmerge_cont = [[False] * n_cols for _ in range(n_rows)]
    # 记录每列最近一次 vMerge restart 的值（用于续行填充）
    vmerge_values = [''] * n_cols

    for r_idx, row in enumerate(table.rows):
        col_cursor = 0
        # 直接遍历行的 XML w:tc 元素，避免 python-docx row.cells 自动扩展合并格导致重复
        tr = row._tr
        for tc in tr.findall(qn('w:tc')):
            # 跳过已被水平合并填充的列（前一个单元格的 gridSpan 填充过）
            while col_cursor < n_cols and grid[r_idx][col_cursor] != '':
                col_cursor += 1
            if col_cursor >= n_cols:
                break

            tcpr = tc.find(qn('w:tcPr'))

            # ── 水平合并（gridSpan） ──
            gridspan = 1
            if tcpr is not None:
                gs_elem = tcpr.find(qn('w:gridSpan'))
                if gs_elem is not None:
                    try:
                        gridspan = int(gs_elem.get(qn('w:val'), '1'))
                    except (ValueError, TypeError):
                        gridspan = 1

            # ── 垂直合并（vMerge） ──
            vmerge_elem = tcpr.find(qn('w:vMerge')) if tcpr is not None else None
            # 提取单元格文本（遍历 w:p/w:r/w:t）
            cell_text_parts = []
            for p_elem in tc.findall(qn('w:p')):
                para_text = ''
                for r_elem in p_elem.findall(qn('w:r')):
                    # ◄ 【优化】识别上标和下标
                    is_sup = False
                    is_sub = False
                    r_pr = r_elem.find(qn('w:rPr'))
                    if r_pr is not None:
                        vert_align = r_pr.find(qn('w:vertAlign'))
                        if vert_align is not None:
                            val = vert_align.get(qn('w:val'), '')
                            if val == 'superscript': is_sup = True
                            elif val == 'subscript': is_sub = True

                    for t_elem in r_elem.findall(qn('w:t')):
                        if t_elem.text:
                            t_text = t_elem.text
                            if is_sup:
                                # 如果是上标且不是以 ^ 开头，补上 ^
                                t_text = f"^{t_text}" if not t_text.startswith('^') else t_text
                            elif is_sub:
                                # 如果是下标且不是以 _ 开头，补上 _
                                t_text = f"_{t_text}" if not t_text.startswith('_') else t_text
                            para_text += t_text
                if para_text.strip():
                    cell_text_parts.append(para_text.strip())
            cell_text = ' '.join(cell_text_parts)

            if vmerge_elem is not None:
                val = vmerge_elem.get(qn('w:val'), '')
                if val == 'restart':
                    # 垂直合并起始行：记录文本
                    vmerge_values[col_cursor] = cell_text
                    actual_text = cell_text
                else:
                    # 垂直合并续行：使用起始行文本，标记为续行
                    actual_text = vmerge_values[col_cursor]
                    for c in range(col_cursor, min(col_cursor + gridspan, n_cols)):
                        is_vmerge_cont[r_idx][c] = True
            else:
                actual_text = cell_text
                vmerge_values[col_cursor] = cell_text  # 更新该列的最新值

            # 填充水平合并的所有列
            for c in range(col_cursor, min(col_cursor + gridspan, n_cols)):
                grid[r_idx][c] = actual_text

            col_cursor += gridspan

    return grid, is_vmerge_cont


def _dedup_row(row: List[str]) -> List[str]:
    """对行进行水平合并去重：相邻相同值只保留一个"""
    result = []
    prev = None
    for cell in row:
        if cell != prev:
            result.append(cell)
            prev = cell
    return result


def _dedup_headers(header_row: List[str]) -> Tuple[List[str], List[int]]:
    """
    对表头行去重，返回 (去重后的表头列表, 对应原始列索引列表)。
    相邻相同的表头（水平合并产生）只保留第一个。
    空表头替换为 column_N。
    """
    unique_headers = []
    kept_indices = []
    prev = None
    for idx, h in enumerate(header_row):
        h_clean = h.strip()
        if h_clean != prev:
            if not h_clean:
                h_clean = f'column_{idx}'
            unique_headers.append(h_clean)
            kept_indices.append(idx)
            prev = h_clean if h_clean else None
        # 相同的相邻值跳过（水平合并重复列）
    return unique_headers, kept_indices


def _extract_msg_name_from_info_row(row_unique: List[str]) -> str:
    """
    从"信息名称行"（行0）的去重单元格列表中提取消息名称。
    过滤掉 INFO_NAME_ROW_NOISE 中所有固定噪声词，取第一个有效值。
    """
    for cell in row_unique:
        cell_clean = cell.strip()
        if cell_clean and cell_clean not in INFO_NAME_ROW_NOISE:
            if len(cell_clean) >= 2 and not cell_clean.isdigit():
                return cell_clean
    return ''


def _extract_name_from_para(para_text: str) -> str:
    """
    从前置段落文本中提取消息/表格名称。
    支持：'表XX 某状态信息'、'表B.1某信道状态'、普通文本
    """
    if not para_text:
        return ''
    text = para_text.strip()

    # 格式1：'表XX 某状态信息'（数字/字母后有空格）
    m = re.match(r'^表[A-Za-z0-9.。\s]*\s+(.+)', text)
    if m:
        name = m.group(1).strip()
        if name and len(name) >= 2:
            return name

    # 格式2：'表B.1某信道状态'（无空格，字母+数字+点后紧跟中文）
    m = re.match(r'^表[A-Za-z0-9.。]+(.+)', text)
    if m:
        name = m.group(1).strip()
        if name and len(name) >= 2:
            return name

    # 格式3：普通文本（非纯序号）
    if len(text) >= 2 and len(text) <= 60:
        if not re.match(r'^[一二三四五六七八九十百千\d\s]+$', text):
            return text

    return ''


def _parse_aggregate_meta(text: str) -> Dict:
    """
    解析聚合式消息元数据字符串。
    支持两种格式：
    - BCRT格式: 'BCRT1-SA0-模式码0x03' → {信源机器码:'BC', 信宿机器码:'1', 子地址:'0', 数据段长度:'3'}
    - RT-SA格式: 'RT13-SA15-8BC' → {信源机器码:'13', 子地址:'15', 数据段长度:'8', 信宿机器码:'BC'}
    """
    meta = {}
    # 格式1: BCRT（信源机器码 BC，信宿机器码 RT后的数字）
    m = re.search(r'(BC)\s*(?:→|->)?\s*RT\s*(\w+)', text, re.IGNORECASE)
    if m:
        meta['信源机器码'] = 'BC'
        meta['信宿机器码'] = m.group(2)

    # 格式2: RT-SA（信源机器码 RT后的数字，信宿机器码 末尾的字母）
    # 示例: RT13-SA15-8BC 或 RT13-SA16-29→BC → 信源机器码=13, 子地址=15/16, 数据段长度=8/29, 信宿机器码=BC
    if '信源机器码' not in meta:
        m = re.search(r'RT(\d+)[-_]?SA(\d+)[-_]?(\d+)(?:→|->)?([A-Za-z]{2})', text, re.IGNORECASE)
        if m:
            meta['信源机器码'] = m.group(1)
            meta['子地址'] = m.group(2)
            meta['数据段长度'] = m.group(3)
            meta['信宿机器码'] = m.group(4).upper()

    # 子地址 SA0（仅在格式2未设置时提取）
    if '子地址' not in meta:
        m = re.search(r'SA\s*(\d+)', text, re.IGNORECASE)
        if m:
            meta['子地址'] = m.group(1)
    # 数据段长度（模式码0x03，仅在格式2未设置时提取）
    if '数据段长度' not in meta:
        m = re.search(r'模式码\s*(0x[0-9A-Fa-f]+|\d+)', text)
        if m:
            val = m.group(1)
            if val.upper().startswith('0X'):
                meta['数据段长度'] = str(int(val, 16))
            else:
                meta['数据段长度'] = val
    return meta


def _is_noise_table(grid: List[List[str]], preceding_para: str, config_field_names: List[str] = None) -> bool:
    """
    判断是否为干扰/无效表格：
    1. 前置段落含干扰词
    2. 是示例/目标格式表（首行就是输出Excel的列名）
    3. 是帧格式说明表（帧头/帧尾）
    
    【不过滤规则】以下表格即使无数据类型关键字也不应被过滤：
    - 端口分配表（含信源系统码+信宿系统码+信息内容）
    - 消息ID表（含消息ID/ID序号+信息内容/ID定义）
    - 匹配用户配置的字段表（表头与配置字段高度重合）
    
    【最后】才检查表头干扰词和整表数据类型关键字。
    配置匹配优先于噪声过滤。
    """
    if any(marker in preceding_para for marker in NOISE_PARA_MARKERS):
        return True
    if not grid:
        return True

    row0_unique = _dedup_row(grid[0])
    row0_text = ' '.join(row0_unique)

    # 检查是否是示例/目标格式表
    row0_set = set(row0_unique)
    if EXAMPLE_TABLE_HEADERS.issubset(row0_set):
        return True

    # 帧格式表
    all_text = ' '.join(cell for row in grid for cell in row)
    if '帧头' in all_text and '帧尾' in all_text:
        return True
    if '帧格式' in preceding_para:
        return True

    # ── 【不过滤】端口分配表 ──────────────────────────────────────────────
    has_port_allocation_features = ('信源系统码' in row0_text and '信宿系统码' in row0_text)
    if has_port_allocation_features:
        return False

    # ── 【不过滤】消息ID表（兼容新旧格式） ───────────────────────────────
    has_old_id = '消息ID' in row0_text and '信息内容' in row0_text
    has_new_id = ('ID序号' in row0_text or ('ID' in row0_text and '序号' in row0_text)) and 'ID定义' in row0_text
    has_generic_id = ('ID' in row0_text) and ('信息内容' in row0_text or '定义' in row0_text or '名称' in row0_text)
    if has_old_id or has_new_id or has_generic_id:
        return False

    # ── 【不过滤】匹配用户配置的字段表 ──────────────────────────────────
    # 如果表头中的字段名与用户配置的字段名高度重合，说明是有效字段表，不过滤
    if config_field_names:
        config_field_set = set(f.strip() for f in config_field_names if f.strip())

        # 辅助函数：计算一行与配置字段的匹配情况
        def _count_config_match(row_cells: List[str]) -> Tuple[int, int]:
            """返回 (匹配数, 非空表头数)"""
            non_empty = [h for h in row_cells if h.strip() and len(h.strip()) >= 2]
            match_count = sum(1 for h in non_empty if h.strip() in config_field_set or
                             any(cf in h or h in cf for cf in config_field_set if len(cf) >= 2))
            return match_count, len(non_empty)

        def _check_match_pass(match_count: int, total: int) -> bool:
            """三档阈值：全部命中 OR >=60% OR >=3个绝对数量"""
            if total == 0:
                return False
            if match_count >= total:
                return True  # 全部命中
            if match_count / total >= 0.60:
                return True  # >=60%
            if match_count >= 3:
                return True  # 绝对数量兜底（避免表头列数多时比例稀释）
            return False

        # 第一档：对 grid[0] 做匹配（适用于类型B/C简单表）
        match_count, total_headers = _count_config_match(row0_unique)
        if _check_match_pass(match_count, total_headers):
            logger.info(f"配置兜底[行0]: 表头与配置字段匹配({match_count}/{total_headers})，不过滤")
            return False

        # 第二档：Type A 表兜底 —— 行0是元数据行（含"通信帧名称/信息名称"），
        # 真正的字段表头在下面某行。向下扫描找候选行。
        #
        # 【兜底策略 - 两阶段扫描】
        #   Phase 1：优先匹配含标准候选词（序号/内容/说明等）的行（兼容老逻辑）
        #   Phase 2：若 Phase1 未命中，扫描所有非元数据行做配置匹配（纯自定义词兜底）
        #           这样即使表头全为自定义词（编号/项目/规格/量程等）也能走配置通道。
        is_type_a_meta_row = any(kw in row0_text for kw in ['信息名称', '通信帧名称'])
        if is_type_a_meta_row:
            # Phase 1：含标准候选词的行（老逻辑，优先匹配）
            for cand_r_idx in range(1, min(8, len(grid))):
                cand_ru = _dedup_row(grid[cand_r_idx])
                cand_rt = ' '.join(cand_ru)
                if any(kw in cand_rt for kw in _TYPE_A_PREFERRED_SCAN_KEYWORDS):
                    c_mc, c_total = _count_config_match(cand_ru)
                    if _check_match_pass(c_mc, c_total):
                        logger.info(f"配置兜底[行{cand_r_idx}][Phase1-含候选词]: TypeA匹配({c_mc}/{c_total})，不过滤")
                        return False

            # Phase 2：纯自定义词兜底 —— 扫描所有非元数据行
            # 跳过：含元数据关键词的行 / 含标准候选词的行（Phase1已处理）/ 无效行
            for cand_r_idx in range(1, min(8, len(grid))):
                cand_ru = _dedup_row(grid[cand_r_idx])
                cand_rt = ' '.join(cand_ru)
                # 跳过含标准候选词的行（Phase1已处理过）
                if any(kw in cand_rt for kw in _TYPE_A_PREFERRED_SCAN_KEYWORDS):
                    continue
                # 跳过含元数据关键词的行
                if any(kw in cand_rt for kw in _TYPE_A_META_KEYWORDS):
                    continue
                # 有效字段行：至少有2个非空、长度>=2、非纯数字的单元格
                valid_cells = [c for c in cand_ru
                               if c.strip() and len(c.strip()) >= 2
                               and not c.strip().lstrip('+-').replace('.', '').isdigit()]
                if len(valid_cells) < 2:
                    continue
                c_mc, c_total = _count_config_match(cand_ru)
                if _check_match_pass(c_mc, c_total):
                    logger.info(f"配置兜底[行{cand_r_idx}][Phase2-纯自定义词]: TypeA匹配({c_mc}/{c_total})，不过滤")
                    return False

    # ── 以下才执行干扰过滤 ────────────────────────────────────────────────
    
    # 检查表头含干扰词
    if any(kw in row0_text for kw in NOISE_HEADER_KEYWORDS):
        return True

    # 检查整表是否含数据类型关键字
    all_upper = all_text.upper()
    has_data_type = any(kw.upper() in all_upper for kw in DATA_TYPE_KEYWORDS)
    return not has_data_type


def _match_target_message_name(grid: List[List[str]], preceding_para: str, target_names: set) -> bool:
    """
    检查表格是否匹配目标消息名称。
    匹配规则：
    1. 表标题（前置段落）是否包含目标名称
    2. 表格第一行（信息名称行）是否包含目标名称
    3. 表格内容中是否包含"信息名称"列，且该列值匹配目标名称
    """
    if not target_names:
        return False
    
    # 规则1：检查前置段落（表标题）
    para_name = _extract_name_from_para(preceding_para)
    if para_name and any(target in para_name for target in target_names):
        return True
    
    # 规则2：检查表格第一行（信息名称行）
    if grid and len(grid) > 0:
        row0_unique = _dedup_row(grid[0])
        msg_name = _extract_msg_name_from_info_row(row0_unique)
        if msg_name and any(target in msg_name for target in target_names):
            return True
    
    # 规则3：检查表格内容中的"信息名称"列
    if grid and len(grid) > 1:
        # 查找包含"信息名称"的列
        headers, kept_indices = _dedup_headers(grid[0])
        info_name_col_idx = None
        for idx, header in enumerate(headers):
            if '信息名称' in header or '通信帧名称' in header:
                info_name_col_idx = kept_indices[idx] if idx < len(kept_indices) else None
                break
        
        # 如果找到"信息名称"列，检查其值
        if info_name_col_idx is not None:
            for row in grid[1:]:  # 跳过表头行
                if info_name_col_idx < len(row):
                    cell_value = row[info_name_col_idx].strip()
                    if cell_value and any(target in cell_value for target in target_names):
                        return True
    
    return False


# ─── 主类 ─────────────────────────────────────────────────────────────────────

class TableDetector:
    """
    协议文档表格检测器（python-docx 版本）。
    完整替换 docx2python，精确处理水平/垂直合并单元格。
    
    修改点：
    1. 支持配置匹配：按"表格类型+字段组合+列角色"传入配置
    2. 处理顺序改为：配置匹配 → 智能识别 → 噪声过滤
    3. 支持 table_type=message_id，兼容旧格式和新格式
    4. 配置兜底：字段表命中配置后必须提取，ID表命中后用于补ID
    """

    def __init__(self, config=None, target_message_names=None):
        # 输出控制：是否丢弃 data_rows 末尾的 CRC 校验字行（默认开启，保持既有行为）
        self.remove_crc_tail = True
        # 保留 keywords 等属性以兼容外部调用
        self.keywords = ['序号', '参数', '内容', '信号名称', '信息内容', '数据类型', '类型',
                         '长度', '单位', '备注', '值域', '信源', '信宿', '消息ID']
        self.noise_markers = ['参见附录']
        self.content_fields = ['参数', '内容', '信号名称', '信息内容', '数据含义', '字段']
        self.header_categories = {
            'sequence': ['序号'],
            'content': ['参数', '内容', '信号名称', '信息内容', '数据含义', '字段'],
            'type': ['数据类型', '类型', '数据格式'],
            'unit': ['单位'],
            'remark': ['备注', '说明', '值域', '数据处理方法'],
            'meta': ['信源', '信宿', '信息内容', '消息ID', '接口名称', '周期', '发起时机', '错误处理']
        }
        # 目标消息名称列表（用于兜底提取）
        self.target_message_names = set(target_message_names) if target_message_names else set()
        
        # 配置兜底：用户传入的表格配置
        self.table_configs = self._parse_config(config)
        
        # 日志记录器
        self.log_records = []
    
    def _parse_config(self, config):
        """
        解析用户配置，提取表格类型和字段组合配置。
        
        支持的配置格式：
        1. 结构化配置（dict with groups / list of dicts）：
           {table_type: 'field_def', required_fields: [...], column_roles: {...}}
        2. 平铺字段名列表（前端直接传来的字段名数组）：
           ['参数', '数据类型', '数据长度', '值域', '单位', '备注']
           自动推断表格类型和角色。
        """
        if not config:
            return []
        
        configs = []
        
        if isinstance(config, list):
            # 检查列表元素类型
            if config and isinstance(config[0], dict):
                # 结构化配置列表
                for item in config:
                    if isinstance(item, dict):
                        configs.append(item)
            elif config and isinstance(config[0], str):
                # ── 平铺字段名列表 → 自动推断表格类型 ──────────────────────
                field_names = [f.strip() for f in config if isinstance(f, str) and f.strip()]
                if field_names:
                    inferred = self._infer_config_from_fields(field_names)
                    configs.extend(inferred)
        elif isinstance(config, dict):
            # 支持配置分组
            groups = config.get('groups', [])
            if groups:
                for group in groups:
                    if isinstance(group, dict):
                        configs.append(group)
            elif 'table_type' in config:
                configs.append(config)
            elif 'fields' in config or 'field_names' in config:
                # dict 格式：字段名列表 + 可选的 ID 表头字段标记
                fields = config.get('fields', config.get('field_names', []))
                id_field_names = config.get('id_field_names', [])
                if fields and isinstance(fields, list):
                    inferred = self._infer_config_from_fields(fields, id_field_names)
                    configs.extend(inferred)
        
        return configs
    
    def _infer_config_from_fields(self, field_names: List[str], id_field_names: List[str] = None) -> List[Dict]:
        """
        根据平铺字段名列表推断表格配置。
        
        推断规则：
        - 如果字段名中包含内容字段（参数/内容/信号名称等）+ 类型字段（数据类型/类型/数据长度等）
          → 推断为 field_def 表
        - 如果用户通过 id_field_names 显式标记了 ID 表字段 → 使用标记创建 message_id 配置
        - 否则回退到严格模式匹配（硬编码集合）
        
        Args:
            field_names: 所有协议字段名列表
            id_field_names: 用户在配置页标记为“ID表头”的字段名列表
        
        Returns:
            推断出的配置列表
        """
        configs = []
        
        # 内容字段候选
        content_candidates = {'参数', '内容', '信号名称', '信息内容', '数据含义', '字段', '名称'}
        # 类型字段候选（包含"长度"类字段，如"数据长度"、"数据长度（字节）"）
        type_candidates = {'数据类型', '类型', '数据格式'}
        length_candidates = {'数据长度', '字节数', '字节', '长度', '数据长度（字节）'}
        # ID 表检测：优先使用用户显式标记，否则严格模式匹配
        id_field_names = id_field_names or []
        has_explicit_id = len(id_field_names) > 0
        
        # 严格回退：只在无显式标记时使用硬编码集合
        id_candidates_strict = {'消息ID', 'ID序号', '消息标识', 'ID定义', 'ID'}
        has_id = has_explicit_id or bool(set(field_names) & id_candidates_strict)
        
        field_set = set(field_names)
        has_content = bool(field_set & content_candidates)
        has_type = bool(field_set & type_candidates)
        has_length = bool(field_set & length_candidates)
        
        # 推断字段定义表
        if has_content and (has_type or has_length):
            configs.append({
                'table_type': 'field_def',
                'required_fields': field_names,
                'column_roles': {},
                '_inferred': True,  # 标记为推断配置
            })
        
        # 推断消息ID表
        if has_id:
            # 列角色推断
            roles = {}
            if has_explicit_id:
                # ── 纯配置驱动：不依赖任何硬编码关键词 ──
                # 按位置分配角色：第1个字段→id_value，第2个→message_name
                for i, f in enumerate(id_field_names):
                    if i == 0:
                        roles['id_value'] = f
                    elif i == 1:
                        roles['message_name'] = f
            else:
                # ── 严格回退：用硬编码关键词（仅当用户未标记 isIdField 时）──
                for f in field_names:
                    if '定义' in f or f == '信息内容' or ('名称' in f and 'ID' not in f):
                        roles['message_name'] = f
                    elif 'ID' in f or 'id' in f.lower():
                        roles['id_value'] = f
            configs.append({
                'table_type': 'message_id',
                'required_fields': field_names,
                'column_roles': roles,
                '_inferred': True,
                '_id_field_names': id_field_names if has_explicit_id else [],
            })
        
        # 如果没有推断出任何类型，创建一个通用字段表配置（宽松匹配）
        if not configs and len(field_names) >= 3:
            configs.append({
                'table_type': 'field_def',
                'required_fields': field_names,
                'column_roles': {},
                '_inferred': True,
                '_lenient': True,  # 宽松匹配模式
            })
        
        return configs
    
    def _match_config(self, grid, headers_text):
        """
        根据配置匹配表格类型。
        匹配规则：
        1. 检查表格类型配置（table_type）
        2. 检查字段组合（配置字段与表头的重合度）
        3. 检查列角色（id_value, message_name等）
        
        匹配策略（按优先级）：
        - 结构化配置：required_fields 中的字段至少有50%在表头中命中
        - 推断配置（_inferred=True）：表头与配置字段的子串匹配比例>=50%
        - 宽松模式（_lenient=True）：只要有内容字段+类型/长度字段命中即可
        
        返回：(匹配的配置, 表格类型) 或 (None, None)
        """
        if not self.table_configs:
            return None, None
        
        # 提取表头字段列表（用于匹配）
        header_fields = [h.strip() for h in headers_text.split() if h.strip()]
        
        for config in self.table_configs:
            table_type = config.get('table_type')
            required_fields = config.get('required_fields', [])
            column_roles = config.get('column_roles', {})
            is_inferred = config.get('_inferred', False)
            is_lenient = config.get('_lenient', False)
            
            if not table_type:
                continue
            
            # ── 字段表配置匹配 ────────────────────────────────────────────
            if table_type == 'field_def':
                matched = self._match_field_def_config(headers_text, header_fields, required_fields, is_lenient)
                if not matched:
                    continue

            # ── ID表配置匹配 ──────────────────────────────────────────────
            elif table_type == 'message_id':
                matched = self._match_message_id_config(headers_text, header_fields, required_fields, column_roles, config)
                if not matched:
                    continue
            
            # ── 其他类型：检查列角色匹配 ──────────────────────────────────
            elif column_roles:
                role_matched = True
                for role, candidates in column_roles.items():
                    if isinstance(candidates, list):
                        found = any(candidate in headers_text for candidate in candidates)
                    else:
                        found = candidates in headers_text
                    if not found:
                        role_matched = False
                        break
                if not role_matched:
                    continue
            
            # 配置匹配成功
            logger.info(f"配置匹配成功: table_type={table_type}, headers_text={headers_text[:50]}...")
            return config, table_type
        
        return None, None
    
    def _match_field_def_config(self, headers_text: str, header_fields: List[str],
                                 required_fields: List[str], is_lenient: bool) -> bool:
        """
        字段定义表配置匹配。
        
        匹配条件：
        1. 必须有内容字段命中（内容/参数/信号名称等）
        2. 必须有类型/长度字段命中（数据类型/类型/数据长度等）
        3. 配置字段与表头的匹配度足够高
        """
        # ── 强制匹配：用户配置了表的全部字段（全字段覆盖） ──
        # 当用户明确给出该表的所有字段名时，即使不含标准关键词
        # （内容/参数/类型/字节），也强制识别为字段定义表，绕过下方硬性要求。
        # 这是"输入所有字段 → 强制识别非标准/自定义协议表"的通道。
        if required_fields and header_fields:
            req_set = set(f.strip() for f in required_fields if f.strip())
            non_empty_hf = [hf for hf in header_fields if hf]
            if len(non_empty_hf) >= 2:
                covered = sum(1 for hf in non_empty_hf if hf in req_set)
                # 表头所有列名都被配置字段覆盖 → 视为用户明确配置，强制匹配
                if covered == len(non_empty_hf):
                    logger.info("配置强制匹配(全字段覆盖): 跳过标准关键词要求，识别为字段定义表")
                    return True

        # 检查内容字段
        content_candidates = ['内容', '参数', '信号名称', '信息内容', '数据含义', '字段']
        has_content_field = any(c in headers_text for c in content_candidates)
        
        # 检查类型/长度字段（扩展：包含"数据长度"类字段）
        type_candidates = ['数据类型', '类型', '数据格式']
        length_candidates = ['数据长度', '字节数', '字节']
        has_type_field = any(c in headers_text for c in type_candidates)
        has_length_field = any(c in headers_text for c in length_candidates)
        
        # 字段表必须同时包含内容字段和类型/长度字段
        if not has_content_field:
            return False
        if not (has_type_field or has_length_field):
            return False
        
        # 宽松模式：只要内容+类型/长度命中就匹配
        if is_lenient:
            return True
        
        # 检查配置字段与表头的匹配度
        if required_fields:
            match_count = 0
            for field in required_fields:
                field_clean = field.strip()
                if not field_clean:
                    continue
                # 精确匹配
                if field_clean in headers_text:
                    match_count += 1
                    continue
                # 子串匹配（表头字段包含配置字段，或反之）
                for hf in header_fields:
                    if field_clean in hf or hf in field_clean:
                        match_count += 1
                        break
            
            match_ratio = match_count / max(len(required_fields), 1)
            # 至少50%的字段匹配
            if match_ratio < 0.5:
                return False
        
        return True
    
    def _match_message_id_config(self, headers_text: str, header_fields: List[str],
                                  required_fields: List[str], column_roles: Dict,
                                  config: Dict = None) -> bool:
        """
        消息ID表配置匹配。
        
        匹配策略：
        1. 有列角色时：所有角色字段必须在表头中命中（精确匹配）
        2. 有 _id_field_names 时：至少 60% 的标记字段必须在表头中命中
        3. 回退：硬编码ID关键词集合（仅未标记 isIdField 时）
        """
        config = config or {}
        
        def _field_in_headers(field_name: str) -> bool:
            """检查字段名是否精确出现在表头中"""
            if not field_name:
                return False
            # 精确匹配：字段名完全等于某个表头
            if field_name in header_fields:
                return True
            # 宽松精确：字段名包含某个表头或反之（但要长度接近）
            for hf in header_fields:
                if len(hf) >= 2 and len(field_name) >= 2:
                    if field_name == hf:
                        return True
                    # 子串匹配但要求长度比 > 0.5（避免单字匹配）
                    if field_name in hf and len(field_name) / len(hf) > 0.5:
                        return True
                    if hf in field_name and len(hf) / len(field_name) > 0.5:
                        return True
            return False
        
        # ── 策略1：列角色全部命中 ────────────────────────────────
        if column_roles:
            all_roles_match = True
            for role, field_name in column_roles.items():
                if isinstance(field_name, str) and field_name:
                    if not _field_in_headers(field_name):
                        all_roles_match = False
                        break
                elif isinstance(field_name, list):
                    if not any(_field_in_headers(f) for f in field_name):
                        all_roles_match = False
                        break
            if all_roles_match:
                return True
        
        # ── 策略2：标记字段至少 60% 命中 ──────────────────────────
        id_field_names = config.get('_id_field_names', [])
        if id_field_names:
            match_count = sum(1 for f in id_field_names if _field_in_headers(f))
            match_ratio = match_count / len(id_field_names)
            if match_ratio >= 0.6:
                return True
        
        # ── 回退：硬编码集合（仅未标记 isIdField 时使用）────────
        id_field_candidates = ['消息ID', 'ID序号', 'ID定义', '消息标识']
        return any(c in headers_text for c in id_field_candidates)
    
    def _log_table_status(self, t_idx, status, reason, table_type=None, msg_name=None, headers=None, preceding=None):
        """记录表格识别状态日志。

        preceding 即表格上方的"表N XXX"标题段落，作为人工定位的主标识
        （列名高度雷同，仅凭表头难以区分具体是哪张表）；headers 作为辅助。
        二者一并写入 log_records，最终进入提取报告。
        """
        clean_headers = [str(h).strip() for h in (headers or []) if str(h).strip()]
        title = (preceding or '').strip()
        self.log_records.append({
            'table_index': t_idx,
            'status': status,
            'reason': reason,
            'table_type': table_type,
            'msg_name': msg_name,
            'headers': clean_headers,
            'preceding_para': title,
        })
        # 表头过长时截断，避免日志行过宽
        hdr_disp = '|'.join(clean_headers[:10])
        if len(hdr_disp) > 80:
            hdr_disp = hdr_disp[:80] + '…'
        # 主标识优先用表标题；缺失时回退到表头
        if title:
            ident = title
            extra = f', 表头: {hdr_disp}' if hdr_disp else ''
        else:
            ident = f'表头: {hdr_disp}' if hdr_disp else '无标题'
            extra = ''
        logger.info(f"Table #{t_idx} [{ident}]: {status} - {reason} (type={table_type}{extra})")

    def extract_tables_from_docx(self, file_path: str) -> List[Dict]:
        """
        主入口：解析 docx 文件，返回结构化表格列表。
        每个表格字典包含：
            index, msg_name, headers, data_rows, meta,
            table_type ('field_def'|'port_allocation'|'message_id'|'bit_def'|'skip'),
            is_auxiliary (bool, 兼容旧接口)
        """
        extracted_tables = []
        try:
            doc = Document(file_path)
        except Exception as e:
            logger.error(f"Failed to open {file_path}: {e}")
            return []

        # 遍历 body 元素，建立段落-表格位置关系
        table_idx = 0
        last_para = ''
        noise_next = False  # 前置段落含干扰词时，标记后续表格为干扰

        for elem in doc.element.body:
            tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag

            if tag == 'p':
                para_text = _get_para_text(elem)
                if para_text:
                    last_para = para_text
                    if any(marker in para_text for marker in NOISE_PARA_MARKERS):
                        noise_next = True
                    # 注意：不重置 noise_next，因为后续连续表格也要跳过
                    # 只有下一个非干扰段落后才重置（见下方逻辑）

            elif tag == 'tbl':
                if table_idx >= len(doc.tables):
                    table_idx += 1
                    continue

                table = doc.tables[table_idx]
                parsed = self._parse_single_table(table, table_idx, last_para, noise_next)
                extracted_tables.append(parsed)

                # 如果当前表格是干扰表，且前置段落也是干扰，保持 noise_next
                # 否则重置（只有明确标记才跳过）
                if not any(marker in last_para for marker in NOISE_PARA_MARKERS):
                    noise_next = False

                table_idx += 1

        return extracted_tables

    def _parse_single_table(self, table, t_idx: int, preceding_para: str, force_skip: bool) -> Dict:
        """
        解析单个表格
        
        修改点：处理顺序改为配置匹配 → 智能识别 → 噪声过滤
        1. 配置匹配：按"表格类型+字段组合+列角色"匹配
        2. 智能识别：端口分配表、消息ID表（兼容新旧格式）、bit位定义表
        3. 噪声过滤：最后判断是否为干扰表
        """
        base = {
            'index': t_idx,
            'msg_name': '',
            'headers': [],
            'data_rows': [],
            'meta': {},
            'table_type': 'skip',
            'is_auxiliary': False,
            'preceding_para': preceding_para,
        }

        try:
            grid, is_vmerge_cont = _build_grid(table)
        except Exception as e:
            logger.warning(f"Table #{t_idx}: grid error: {e}")
            return base

        if not grid or len(grid) < 1:
            return base

        # 提前收集配置字段名（供配置匹配、噪声过滤、字段表解析共用）
        config_field_names = []
        for cfg in self.table_configs:
            config_field_names.extend(cfg.get('required_fields', []))
        config_field_names = [f.strip() for f in config_field_names if f.strip()]

        # 获取行0去重内容
        row0_unique = _dedup_row(grid[0])
        row0_text = ' '.join(row0_unique)
        
        # ── 修改点1：配置匹配（最高优先级） ─────────────────────────────────────
        # 【例外】示例/目标格式表始终过滤，即使配置匹配也不提取
        row0_set = set(h.strip() for h in row0_unique if h.strip())
        if EXAMPLE_TABLE_HEADERS.issubset(row0_set):
            self._log_table_status(t_idx, '过滤', '示例/目标格式表', 'skip', headers=row0_unique, preceding=preceding_para)
            return base
        
        config, matched_table_type = self._match_config(grid, row0_text)
        if config and matched_table_type:
            # 配置匹配成功，强制提取
            self._log_table_status(t_idx, '配置匹配', f'命中配置: {matched_table_type}', matched_table_type, headers=row0_unique, preceding=preceding_para)
            
            if matched_table_type == 'field_def':
                return self._parse_field_def_table(grid, is_vmerge_cont, t_idx, preceding_para, config_field_names)
            elif matched_table_type == 'message_id':
                config_roles = config.get('column_roles', {})
                return self._parse_message_id_table(grid, t_idx, preceding_para, config_roles)
            elif matched_table_type == 'port_allocation':
                return self._parse_port_allocation(grid, t_idx, preceding_para)
            elif matched_table_type == 'bit_def':
                return self._parse_bit_def_table(grid, t_idx, preceding_para)
        
        # ── 修改点2：智能识别（中等优先级） ─────────────────────────────────────
        
        # 端口分配表识别（含信源系统码+信宿系统码）
        if '信源系统码' in row0_text and '信宿系统码' in row0_text:
            self._log_table_status(t_idx, '智能识别', '端口分配表', 'port_allocation', headers=row0_unique, preceding=preceding_para)
            return self._parse_port_allocation(grid, t_idx, preceding_para)

        # 消息ID表识别（兼容新旧格式 + 配置驱动）
        # 旧格式：消息ID + 信息内容
        # 新格式：ID序号 + ID定义 + 是否有数据（ID序号为id_value，ID定义为message_name）
        # 配置驱动：column_roles 指定 id_value 和 message_name 对应的列名
        row0_text_lower = row0_text.lower()
        has_message_id = '消息ID' in row0_text
        has_info_content = '信息内容' in row0_text
        has_id_seq = 'ID序号' in row0_text or ('序号' in row0_text and 'id' in row0_text_lower)
        has_id_def = 'ID定义' in row0_text
        # 智能识别回退（仅当配置未匹配时到达此处）：含“id”和“定义/名称”等组合
        has_generic_id = 'id' in row0_text_lower and ('定义' in row0_text or '名称' in row0_text)
        
        # 从配置中获取ID表列角色提示
        config_id_roles = {}
        if config and matched_table_type == 'message_id':
            config_id_roles = config.get('column_roles', {})
        
        is_id_table = (has_message_id and has_info_content) or \
                      (has_id_seq and has_id_def) or \
                      has_generic_id or \
                      bool(config_id_roles)
        
        if is_id_table:
            unique_cols, _ = _dedup_headers(grid[0])
            if len(unique_cols) <= 8:
                self._log_table_status(t_idx, '智能识别', '消息ID表', 'message_id', headers=row0_unique, preceding=preceding_para)
                return self._parse_message_id_table(grid, t_idx, preceding_para, config_id_roles)

        # bit位定义表识别（含位号+状态参数）
        if ('位号' in row0_text or '位号' in ' '.join(_dedup_row(grid[1])) if len(grid) > 1 else False) \
                and '状态参数' in row0_text:
            self._log_table_status(t_idx, '智能识别', 'bit位定义表', 'bit_def', headers=row0_unique, preceding=preceding_para)
            return self._parse_bit_def_table(grid, t_idx, preceding_para)
        # 也匹配只有"位号"列的情况
        if '位号' in row0_text and any(kw in row0_text for kw in ['状态参数', '取值说明']):
            self._log_table_status(t_idx, '智能识别', 'bit位定义表', 'bit_def', headers=row0_unique, preceding=preceding_para)
            return self._parse_bit_def_table(grid, t_idx, preceding_para)
        
        # ── 修改点3：噪声过滤（最低优先级） ─────────────────────────────────────
        
        # 强制跳过（除非匹配目标名称）
        if force_skip:
            if not _match_target_message_name(grid, preceding_para, self.target_message_names):
                self._log_table_status(t_idx, '过滤', '前置段落干扰词', 'skip', headers=row0_unique, preceding=preceding_para)
                return base

        # 判断是否为干扰/无效表格
        # 传入配置字段名列表，让噪声过滤函数知道哪些表头是用户期望的
        if _is_noise_table(grid, preceding_para, config_field_names):
            # 检查是否匹配目标名称，如果匹配则强制提取
            if not _match_target_message_name(grid, preceding_para, self.target_message_names):
                self._log_table_status(t_idx, '过滤', '噪声表', 'skip', headers=row0_unique, preceding=preceding_para)
                return base
            else:
                self._log_table_status(t_idx, '强制提取', '匹配目标名称，跳过噪声过滤', 'field_def', headers=row0_unique, preceding=preceding_para)
                logger.info(f"Table #{t_idx}: 匹配目标名称，强制跳过干扰表判断")
        
        # ── 字段定义表（最后处理） ──────────────────────────────────────────────
        self._log_table_status(t_idx, '智能识别', '字段定义表', 'field_def', headers=row0_unique, preceding=preceding_para)
        return self._parse_field_def_table(grid, is_vmerge_cont, t_idx, preceding_para, config_field_names)

    # ── 端口分配表 ────────────────────────────────────────────────────────────

    def _parse_port_allocation(self, grid: List[List[str]], t_idx: int, preceding_para: str) -> Dict:
        headers, kept_indices = _dedup_headers(grid[0])
        data_rows = self._extract_data_rows(grid, headers, kept_indices, start_row=1)
        return {
            'index': t_idx,
            'msg_name': '端口分配表',
            'headers': headers,
            'data_rows': data_rows,
            'meta': {},
            'table_type': 'port_allocation',
            'is_auxiliary': True,
            'preceding_para': preceding_para,
        }

    # ── 消息ID映射表 ──────────────────────────────────────────────────────────

    def _parse_message_id_table(self, grid: List[List[str]], t_idx: int, preceding_para: str,
                                  config_column_roles: Dict = None) -> Dict:
        """
        解析消息ID表，兼容新旧格式：
        - 旧格式：消息ID + 信息内容
        - 新格式：ID序号 + ID定义 + 是否有数据
        - 配置驱动：按 column_roles 指定列角色
        
        列角色确定优先级：
        1. 配置中的 column_roles（最高优先级）
        2. 智能识别（消息ID/ID序号 → id_value，信息内容/ID定义 → message_name）
        """
        headers, kept_indices = _dedup_headers(grid[0])
        data_rows = self._extract_data_rows(grid, headers, kept_indices, start_row=1)
        
        meta = {}
        
        # ── 优先使用配置指定的列角色 ──────────────────────────────────────
        if config_column_roles:
            id_col = config_column_roles.get('id_value', '')
            name_col = config_column_roles.get('message_name', '')
            if id_col:
                meta['id_column'] = id_col
            if name_col:
                meta['name_column'] = name_col
        
        # ── 如果配置未指定，先用关键词尝试，再回退到数据模式推断 ───
        if 'id_column' not in meta or 'name_column' not in meta:
            # 第一轮：关键词匹配（向后兼容）
            for header in headers:
                if 'name_column' not in meta:
                    if '定义' in header or header == '信息内容' or ('名称' in header and 'id' not in header.lower()):
                        meta['name_column'] = header
                if 'id_column' not in meta:
                    if ('id' in header.lower()) and '定义' not in header and '说明' not in header:
                        meta['id_column'] = header
        
        # 第二轮：数据模式推断（关键词全都不匹配时，分析实际数据）
        if ('id_column' not in meta or 'name_column' not in meta) and data_rows:
            import re as _re
            sample_rows = data_rows[:min(5, len(data_rows))]
            for i, header in enumerate(headers):
                if header in meta.values():
                    continue
                vals = [r.get(header, '') for r in sample_rows if r.get(header)]
                if not vals:
                    continue
                # 超过一半的值像 ID（0x开头或纯数字）→ id_column
                hex_count = sum(1 for v in vals if _re.match(r'^0x[0-9A-Fa-f]+$', str(v).strip()))
                num_count = sum(1 for v in vals if _re.match(r'^\d+$', str(v).strip()))
                if hex_count > len(vals) / 2 or num_count > len(vals) / 2:
                    if 'id_column' not in meta:
                        meta['id_column'] = header
                else:
                    if 'name_column' not in meta:
                        meta['name_column'] = header
        
        return {
            'index': t_idx,
            'msg_name': '消息ID表',
            'headers': headers,
            'data_rows': data_rows,
            'meta': meta,
            'table_type': 'message_id',
            'is_auxiliary': True,
            'preceding_para': preceding_para,
        }

    # ── bit位定义表 ───────────────────────────────────────────────────────────

    def _parse_bit_def_table(self, grid: List[List[str]], t_idx: int, preceding_para: str) -> Dict:
        headers, kept_indices = _dedup_headers(grid[0])
        data_rows = self._extract_data_rows(grid, headers, kept_indices, start_row=1)
        msg_name = _extract_name_from_para(preceding_para)
        return {
            'index': t_idx,
            'msg_name': msg_name,
            'headers': headers,
            'data_rows': data_rows,
            'meta': {},
            'table_type': 'bit_def',
            'is_auxiliary': True,
            'preceding_para': preceding_para,
        }

    # ── 字段定义表（A/B/C 三种类型 + 聚合式） ─────────────────────────────────

    def _parse_field_def_table(self, grid: List[List[str]], is_vmerge_cont: List[List[bool]],
                                t_idx: int, preceding_para: str,
                                config_field_names: List[str] = None) -> Dict:
        n_rows = len(grid)
        cfg_set = set(f.strip() for f in (config_field_names or []) if f.strip())
        skip_result = {
            'index': t_idx, 'msg_name': '', 'headers': [], 'data_rows': [],
            'meta': {}, 'table_type': 'skip', 'is_auxiliary': False,
            'preceding_para': preceding_para,
        }

        # ── 判断是否有"信息名称行"（行0 含 信息名称/通信帧名称） ──────────────
        row0_unique = _dedup_row(grid[0])
        row0_text = ' '.join(row0_unique)
        # 聚合式表格识别：行0含特定关键词 或 前置段落含"聚合式"
        has_info_name_row = any(kw in row0_text for kw in ['信息名称', '通信帧名称'])
        is_aggregate_table = '聚合式' in preceding_para

        header_row_idx = -1
        msg_name = ''
        meta = {}

        if has_info_name_row or is_aggregate_table:
            # 类型A / 聚合式：行0 是信息名称行 或 前置段落标记为聚合式
            if has_info_name_row:
                msg_name = _extract_msg_name_from_info_row(row0_unique)
            else:
                # 来自聚合式前置段落，消息名称从前置段落提取
                msg_name = _extract_name_from_para(preceding_para)

            # ◄ 【新增】名称优先级：优先从横向元数据结构中提取
            # 聚合式表格的元数据是横向排列的，且由于合并单元格，实际列排列为：
            # 列0=键1（如"数据项名称"），列1=键1_重复（合并），列2=值1（如"计算结果"），列3=值1_重复（合并）...
            # 或更简单地说：找第一个非空的非重复值（不是"数据项名称"等元数据键）
            if is_aggregate_table and n_rows > 0:
                row0 = grid[0]
                if row0 and len(row0) >= 2:
                    # 在行0中查找表名（通常是第一个不是元数据键的值）
                    metadata_keys = {'数据项名称', '消息名称', '表名', '通信名称', '数据流名称'}
                    for col_idx, cell_val in enumerate(row0):
                        cell_clean = cell_val.strip()
                        # 跳过元数据键和空值
                        if cell_clean and cell_clean not in metadata_keys and cell_clean not in INFO_NAME_ROW_NOISE:
                            # 检查这是否看起来像表名（不含数据类型关键字）
                            if not any(kw in cell_clean for kw in ['字节', 'UINT', 'INT', 'FLOAT', '序号', '内容']):
                                msg_name = cell_clean  # 【最高优先级】从横向元数据取表名
                                break
            
            # 备选：从表内第一列纵向元数据提取（如果上面的逻辑不适用）
            if not msg_name and is_aggregate_table and n_rows > 1:
                for check_row in range(1, min(4, n_rows)):  # 检查行1-3
                    first_col_val = grid[check_row][0].strip()  # 第一列
                    if first_col_val and first_col_val not in INFO_NAME_ROW_NOISE:
                        # 检查这是否是有效的表名（不是元数据关键词）
                        if not any(kw in first_col_val for kw in ['发起时机', '错误处理', '传输周期', '按实际操作流程', '序号', '检查结果']):
                            # 再检查这一行是否看起来像表名而不是数据行
                            row_content = ' '.join(_dedup_row(grid[check_row]))
                            if not any(kw in row_content for kw in ['字节', '字节数', 'UINT', 'INT', 'FLOAT']):
                                msg_name = first_col_val  # 【次优先级】
                                break

            # 从行1开始找真正的列名行（含数据类型/内容等关键字）
            # 对于聚合式表格，表头可能在行1-7之间的任何位置
            start_search = 1 if has_info_name_row else 0
            for r_idx in range(start_search, min(8, n_rows)):
                ru = _dedup_row(grid[r_idx])
                rt = ' '.join(ru)
                has_type = any(kw in rt for kw in ['数据类型', '类型', '数据格式', '字节', '字节数', '长度', '数据长度'])
                # 对于聚合式表格，需要更精确地识别内容列（排除元数据行）
                has_content = any(kw in rt for kw in ['内容', '参数', '信号名称', '字段', '数据含义', '名称'])
                # 添加聚合式特有的表头标记
                has_aggregate_marker = any(kw in rt for kw in ['计算结果', '消息类型', '消息序号']) and has_type
                has_seq = '序号' in rt
                
                # 排除元数据行（聚合式表格的元数据区特征：只有特定的关键词，没有"内容"列）
                # 元数据行特征：包含元数据关键词且行只有2-3个唯一的非空值
                is_metadata_row_pattern = False
                if has_info_name_row or is_aggregate_table:
                    non_empty_cells = [c for c in ru if c and c not in ('—', '-')]
                    metadata_keywords = {'信息名称', '信息标识', '信源、信宿', '信源、信目', '传输周期', '发起时机', '错误处理'}
                    # 如果行包含元数据关键词且无"内容"和"序号"，则是元数据行
                    has_metadata_kw = any(kw in rt for kw in metadata_keywords)
                    is_metadata_row_pattern = has_metadata_kw and not has_content and not has_seq
                    
                    # 对于聚合式表格：如果只有"计算结果"/"消息类型"等特殊列但没有"序号"或真正的"内容"列
                    # 说明这不是字段定义表的表头，而是元数据行的变体
                    if not is_metadata_row_pattern and r_idx < 3 and is_aggregate_table:
                        # 行0-2：检查是否缺少"序号"且只有特殊关键词
                        has_special_marker = any(kw in rt for kw in ['计算结果', '消息类型', '消息序号', '信源、信宿', '信源、信目'])
                        if has_special_marker and not has_seq and not (has_content and has_type):
                            # 这是元数据行，跳过
                            is_metadata_row_pattern = True
                
                if is_metadata_row_pattern:
                    continue

                # 判断是否是表头行
                is_header_by_kw = has_type or (has_content and has_seq) or \
                                  (has_content and len(ru) >= 3) or has_aggregate_marker

                # ── 配置兜底：关键词未命中时，用配置字段完全覆盖判定 ──
                # 用户配齐该表字段时，即使列名全是自定义词（编号/项目/规格…）
                # 也能强制认定为表头行。条件：该行所有列名都在配置字段中，
                # 或（命中≥3 且 覆盖率≥60%）。
                is_header_by_cfg = False
                if not is_header_by_kw and cfg_set:
                    non_empty = [c for c in ru if c and c.strip()]
                    if len(non_empty) >= 2:
                        covered = sum(1 for c in non_empty if c.strip() in cfg_set)
                        if covered == len(non_empty) or \
                           (covered >= 3 and covered / len(non_empty) >= 0.6):
                            is_header_by_cfg = True
                            logger.info(f"配置兜底[行{r_idx}]: 表头全部由配置字段覆盖，强制认定为字段表头")

                if is_header_by_kw or is_header_by_cfg:
                    header_row_idx = r_idx
                    break

            # 聚合式：收集元数据区（行1到表头行之间）
            # 支持两种格式：
            # 1. 纵向：行1列0=键，后续行处理
            # 2. 横向：行0-N中，列0=键1，列1=值1，列2=键2，列3=值2...
            if has_info_name_row and header_row_idx >= 2:
                for r_idx in range(1, header_row_idx):
                    row = _dedup_row(grid[r_idx])
                    row_text = ' '.join(row)
                    
                    # ◄ 新增：横向元数据提取（键值对横向排列）
                    # 特别处理这一行，假设列0=键1，列1=值1，列2=键2，列3=值2...
                    for col_idx in range(0, len(row) - 1, 2):
                        key_cell = row[col_idx].strip()
                        val_cell = row[col_idx + 1].strip() if col_idx + 1 < len(row) else ''
                        
                        if key_cell and val_cell and key_cell not in ('—', '-', ''):
                            # 标准化元数据键
                            if key_cell in ['信源、信宿', '信源、信目']:
                                meta['信源、信宿'] = val_cell
                            elif '发起时机' in key_cell:
                                meta['发起时机'] = val_cell
                            elif '发送周期' in key_cell:
                                meta['发送周期'] = val_cell
                            elif '错误处理' in key_cell:
                                meta['错误处理'] = val_cell
                            elif '备注' in key_cell:
                                meta['备注'] = val_cell
                    
                    # ◄ 原有逻辑：处理特殊格式（BCRT、SA等）
                    if any(kw in row_text for kw in ['BCRT', 'SA', '模式码', '→', 'BC']):
                        meta.update(_parse_aggregate_meta(row_text))
                    
                    # ◄ 原有逻辑：提取信源信宿（IP地址格式）
                    if '信源、信宿' in row_text or '信源、信目' in row_text:
                        for cell in row:
                            if '→' in cell or ':' in cell:
                                meta['信源、信宿'] = cell
        else:
            # 类型B / C：行0 就是列名行
            row0_text = ' '.join(_dedup_row(grid[0]))
            # 扩展类型字段检测：包含"数据长度"类字段（如"数据长度（字节）"）
            has_type = any(kw in row0_text for kw in ['数据类型', '类型', '数据格式', '数据长度', '字节数'])
            has_content = any(kw in row0_text for kw in ['内容', '参数', '信号名称', '字段', '数据含义', '名称'])
            if has_type or has_content:
                header_row_idx = 0
                # 消息名称来自前置段落标题
                msg_name = _extract_name_from_para(preceding_para)
            else:
                # 配置兜底：行0 列名全部在配置字段中 → 强制认定为字段表头
                header_row_idx = -1
                if cfg_set:
                    row0_cells = [c for c in _dedup_row(grid[0]) if c and c.strip()]
                    if len(row0_cells) >= 2:
                        covered = sum(1 for c in row0_cells if c.strip() in cfg_set)
                        if covered == len(row0_cells) or \
                           (covered >= 3 and covered / len(row0_cells) >= 0.6):
                            header_row_idx = 0
                            msg_name = _extract_name_from_para(preceding_para)
                            logger.info("配置兜底[B/C]: 行0 列名全部由配置字段覆盖，强制认定为字段表头")
                if header_row_idx == -1:
                    return skip_result

        if header_row_idx == -1:
            return skip_result

        # ── 提取列名（去重） ──────────────────────────────────────────────────
        headers, kept_indices = _dedup_headers(grid[header_row_idx])

        # ── 提取数据行 ─────────────────────────────────────────────────────────
        data_rows = self._extract_data_rows(grid, headers, kept_indices,
                                             start_row=header_row_idx + 1,
                                             is_vmerge_cont=is_vmerge_cont)

        if not data_rows:
            return skip_result

        # ── 判断是否是辅助表（端口/ID 已经提前识别，这里只标记字段定义表） ──
        is_auxiliary = False

        return {
            'index': t_idx,
            'msg_name': msg_name,
            'headers': headers,
            'data_rows': data_rows,
            'meta': meta,
            'table_type': 'field_def',
            'is_auxiliary': is_auxiliary,
            'preceding_para': preceding_para,
        }

    # ── 通用数据行提取 ─────────────────────────────────────────────────────────

    def _extract_data_rows(self, grid: List[List[str]], headers: List[str],
                           kept_indices: List[int], start_row: int,
                           is_vmerge_cont: Optional[List[List[bool]]] = None) -> List[Dict]:
        """
        从 grid 的 start_row 开始提取数据行，按 headers/kept_indices 映射列。
        过滤空行、注释行、无效元数据行（只有名称和内容但无其他数据的行）。
        """
        data_rows = []
        # 定义内容字段名称集合（这些字段用于存放数据名称或描述）
        content_field_names = {'名称', '内容', '参数', '信号名称', '字段', '数据含义', '参数名称', '数据项名称', '代号', '描述'}
        # 聚合式表格元数据行关键词（这些作为"内容"出现时，表示是元数据行而非数据行）
        # 注意：只匹配精确的元数据行标记，避免误匹配包含这些词的字段名（如"消息序号"不应被当作"序号"元数据）
        metadata_row_keywords = {
            '聚合式的信息流表征示意', '信息名称行', '信息标识行', '信源、信宿', '信源、信目',
            '传输周期', '发起时机', '错误处理', '^序号$',  # 用正则表达式匹配单独的"序号"
            '检查结果', '非周期', '按实际操作流程'
        }

        # 内容字段（数据含义/名称等）所在的 grid 列，用于判断垂直合并续行
        content_col_idx = None
        for h, col_idx in zip(headers, kept_indices):
            if h in content_field_names:
                content_col_idx = col_idx
                break

        for r_idx in range(start_row, len(grid)):
            row = grid[r_idx]

            row_dict = {}
            for h_idx, (h, col_idx) in enumerate(zip(headers, kept_indices)):
                val = row[col_idx] if col_idx < len(row) else ''
                row_dict[h] = val.strip()

            # 过滤纯空行
            row_all = ''.join(row_dict.values())
            if not row_all.strip():
                continue

            # ◄ 垂直合并续行合并：
            # 若内容字段（如“数据含义”）所在单元格是 docx 垂直合并的续行，说明本行与上一
            # 数据项属于同一个合并单元格（例如“设备状态”跨多行，而“备注”列 D1~D8 各自独立）。
            # 此时把本行中“与上一行不同的非空值”追加到上一行（换行连接），而非新增数据行。
            is_continuation = (
                is_vmerge_cont is not None
                and content_col_idx is not None
                and r_idx < len(is_vmerge_cont)
                and content_col_idx < len(is_vmerge_cont[r_idx])
                and is_vmerge_cont[r_idx][content_col_idx]
            )
            if is_continuation and data_rows:
                prev = data_rows[-1]
                for header, value in row_dict.items():
                    v = (value or '').strip()
                    if not v or v in ('—', '-'):
                        continue
                    prev_v = (prev.get(header) or '').strip()
                    if v == prev_v:
                        continue
                    prev[header] = (prev_v + '\n' + v) if prev_v else v
                continue

            # 过滤注释行
            if any(row_all.startswith(p) for p in ['注：', '注:', '说明：', '说明:']):
                continue

            # 过滤噪声（参见附录等）
            if any(m in row_all for m in self.noise_markers):
                continue

            # 至少要有1个有实际意义的字段
            non_empty = [v for v in row_dict.values() if v and v not in ('—', '-', '序号', '')]
            if not non_empty:
                continue

            # ◄ 元数据行检测（聚合式表格中的元数据区）
            # 如果行的内容字段（名称、内容等）包含元数据关键词，则过滤掉
            is_metadata_row = False
            for header, value in row_dict.items():
                if header in content_field_names and value:
                    value_str = str(value).strip()
                    # 检查该内容字段是否包含元数据关键词
                    for keyword in metadata_row_keywords:
                        if keyword.startswith('^') and keyword.endswith('$'):
                            # 正则表达式：精确匹配
                            pattern = keyword[1:-1]
                            if re.fullmatch(pattern, value_str):
                                is_metadata_row = True
                                break
                        elif keyword in value_str:
                            # 子串匹配
                            is_metadata_row = True
                            break
                if is_metadata_row:
                    break
            
            if is_metadata_row:
                continue

            # ◄ 核心过滤：如果只有名称/内容字段有值，但其他数据字段全空，则过滤掉
            # 这些通常是错误包含的元数据行，如"聚合式的信息流表征示意"
            has_non_content_data = False
            for header, value in row_dict.items():
                # 跳过内容字段（名称、内容等描述字段）
                if header in content_field_names:
                    continue
                # 检查其他字段是否有实际数据
                if value and str(value).strip() and str(value).strip() not in ('—', '-', ''):
                    has_non_content_data = True
                    break
            
            # 如果没有任何非内容字段的数据，说明这是无效行，过滤掉
            if not has_non_content_data:
                continue

            data_rows.append(row_dict)

        # ◄ 末尾 CRC 校验字过滤（可由输出控制开关关闭）：
        # 若最后一个有效数据行的内容字段包含 "CRC校验"/"CRC检验"（CRC 大小写不敏感），
        # 则丢弃该行（这类校验字段通常不是协议有效数据项），可覆盖 CRC校验字/校验码/校验 等写法。
        # 仅作用于最后一项：若 CRC 行后面仍有有效数据行，则保留（见 content 居中的情况）。
        if self.remove_crc_tail:
            crc_keywords = ('crc校验', 'crc检验')
            if data_rows:
                last_row = data_rows[-1]
                for header, value in last_row.items():
                    if header in content_field_names and value:
                        v = str(value).strip().lower()
                        if any(kw in v for kw in crc_keywords):
                            data_rows.pop()
                            break

        return data_rows


# ─── DocumentParser ──────────────────────────────────────────────────────────

class DocumentParser:
    def __init__(self, config=None, target_message_names=None):
        self.detector = TableDetector(config, target_message_names)
        try:
            from .table_linker import TableLinker
        except ImportError:
            try:
                from backend.services.table_linker import TableLinker
            except ImportError:
                import sys, os
                sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))
                from backend.services.table_linker import TableLinker
        self.linker = TableLinker()

    def parse(self, path: str, options: Dict = None) -> Dict:

        # 应用输出控制选项
        options = options or {}
        if 'remove_crc_tail' in options:
            self.detector.remove_crc_tail = bool(options['remove_crc_tail'])

        raw_tables = self.detector.extract_tables_from_docx(path)
        linked_tables = self.linker.link_tables(raw_tables)
        
        # ◄ 【新增】输出识别结果为JSON
        self._output_recognition_results(raw_tables, linked_tables, path)
        
        return {'tables': linked_tables, 'tables_count': len(linked_tables)}
    
    def _output_recognition_results(self, raw_tables: List[Dict], linked_tables: List[Dict], doc_path: str):
        """输出表格识别结果为JSON文件（包含所有表格类型和分类日志，用于调试和验证）"""
        import json
        import os
        from datetime import datetime
        
        output_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(doc_path))), 'table_recognition_results')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            print(f"-> 已创建输出目录: {output_dir}")
        
        basename = os.path.basename(doc_path)
        
        # ── 1. 分类日志（每个表格的识别决策过程）──────────────────────
        log_file = os.path.join(output_dir, '1_classification_log.json')
        log_data = {
            'file': basename,
            'timestamp': datetime.now().isoformat(),
            'total_tables_scanned': len(self.detector.log_records),
            'decisions': self.detector.log_records
        }
        try:
            with open(log_file, 'w', encoding='utf-8') as f:
                json.dump(log_data, f, ensure_ascii=False, indent=2)
            print(f"[INFO] 分类日志已保存: {log_file}")
        except Exception as e:
            print(f"[ERROR] 保存分类日志失败: {e}")
        
        # ── 2. 原始识别结果（所有表格，含 ID/端口/bit/field_def/skip）────
        raw_file = os.path.join(output_dir, '2_raw_tables.json')
        raw_data = {
            'file': basename,
            'timestamp': datetime.now().isoformat(),
            'total_tables': len(raw_tables),
            'tables': []
        }
        for table in raw_tables:
            table_info = {
                'index': table.get('index'),
                'msg_name': table.get('msg_name', ''),
                'table_type': table.get('table_type', 'unknown'),
                'is_auxiliary': table.get('is_auxiliary', False),
                'headers': table.get('headers', []),
                'data_rows_count': len(table.get('data_rows', [])),
                'meta': table.get('meta', {}),
                'data_rows': [
                    {k: (str(v)[:120] if v else '') for k, v in row.items() if not str(k).startswith('_')}
                    for row in table.get('data_rows', [])[:15]
                ]
            }
            raw_data['tables'].append(table_info)
        try:
            with open(raw_file, 'w', encoding='utf-8') as f:
                json.dump(raw_data, f, ensure_ascii=False, indent=2)
            print(f"[INFO] 原始表格已保存: {raw_file}")
        except Exception as e:
            print(f"[ERROR] 保存原始表格失败: {e}")
        
        # ── 3. 关联后结果（仅 field_def，含注入的元数据和 bit 子行）────
        linked_file = os.path.join(output_dir, '3_linked_tables.json')
        linked_data = {
            'file': basename,
            'timestamp': datetime.now().isoformat(),
            'total_field_def_tables': len(linked_tables),
            'tables': []
        }
        for table in linked_tables:
            table_info = {
                'index': table.get('index'),
                'msg_name': table.get('msg_name', ''),
                'table_type': table.get('table_type', ''),
                'headers': table.get('headers', []),
                'data_rows_count': len(table.get('data_rows', [])),
                'meta': table.get('meta', {}),
                'meta_sources': table.get('meta_sources', {}),
                'data_rows': [
                    {k: (str(v)[:120] if v else '') for k, v in row.items() if not str(k).startswith('_')}
                    for row in table.get('data_rows', [])[:15]
                ]
            }
            linked_data['tables'].append(table_info)
        try:
            with open(linked_file, 'w', encoding='utf-8') as f:
                json.dump(linked_data, f, ensure_ascii=False, indent=2)
            print(f"[INFO] 关联表格已保存: {linked_file}")
        except Exception as e:
            print(f"[ERROR] 保存关联表格失败: {e}")
        
        # ── 5. latest_recognition.json（综合摘要，兼容旧格式）────────
        latest_file = os.path.join(output_dir, 'latest_recognition.json')
        latest_data = {
            'file': basename,
            'timestamp': datetime.now().isoformat(),
            'total_tables': len(linked_tables),
            'tables': linked_data['tables']
        }
        try:
            with open(latest_file, 'w', encoding='utf-8') as f:
                json.dump(latest_data, f, ensure_ascii=False, indent=2)
            print(f"[INFO] 综合摘要已保存: {latest_file}")
        except Exception as e:
            print(f"[ERROR] 保存综合摘要失败: {e}")
