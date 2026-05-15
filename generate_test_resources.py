#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试文档生成器
创建用于字段配置功能测试的复杂文档
"""

import json
import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document():
    """创建字段配置测试文档"""
    doc = Document()
    
    # 标题
    title = doc.add_heading('字段配置功能测试文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 文档说明
    doc.add_paragraph('本文档用于测试字段配置功能，包含多种字段格式和表格结构')
    doc.add_paragraph('测试目标：验证系统对不同字段名称的识别和映射能力')
    
    # 第一个表格 - 标准字段格式
    doc.add_heading('1. 标准协议字段表', level=1)
    table1 = doc.add_table(rows=1, cols=5)
    table1.style = 'Table Grid'
    
    # 表头
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '参数名称'
    hdr_cells[2].text = '数据类型说明'
    hdr_cells[3].text = '单位'
    hdr_cells[4].text = '备注信息'
    
    # 数据行
    data_rows = [
        ['1', '温度传感器', 'float32', '摄氏度', '主传感器读数'],
        ['2', '压力计', 'int16', '帕斯卡', '系统压力监测'],
        ['3', '电压表', 'float64', '伏特', '电源电压检测']
    ]
    
    for row_data in data_rows:
        row_cells = table1.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    doc.add_paragraph()
    
    # 第二个表格 - 非标准字段格式
    doc.add_heading('2. 非标准字段格式表', level=1)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    
    # 表头（使用非标准名称）
    hdr_cells = table2.rows[0].cells
    hdr_cells[0].text = '编号'
    hdr_cells[1].text = '字段名'
    hdr_cells[2].text = '类型描述'
    hdr_cells[3].text = '计量单位'
    
    # 数据行
    data_rows2 = [
        ['A001', '湿度检测器', '整型数值', '百分比'],
        ['A002', '光照强度', '浮点数', '勒克斯'],
        ['A003', '风速计', '双精度', '米/秒']
    ]
    
    for row_data in data_rows2:
        row_cells = table2.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    doc.add_paragraph()
    
    # 第三个表格 - 混合字段格式
    doc.add_heading('3. 混合字段格式表', level=1)
    table3 = doc.add_table(rows=1, cols=6)
    table3.style = 'Table Grid'
    
    # 表头（混合标准和非标准）
    hdr_cells = table3.rows[0].cells
    hdr_cells[0].text = 'ID'
    hdr_cells[1].text = '参数'
    hdr_cells[2].text = '格式'
    hdr_cells[3].text = '单位'
    hdr_cells[4].text = '范围'
    hdr_cells[5].text = '说明'
    
    # 数据行
    data_rows3 = [
        ['001', 'CPU温度', 'float', '℃', '0-100', '处理器核心温度'],
        ['002', '内存使用率', 'int', '%', '0-100', 'RAM占用百分比'],
        ['003', '网络延迟', 'double', 'ms', '0-1000', 'ping响应时间']
    ]
    
    for row_data in data_rows3:
        row_cells = table3.add_row().cells
        for i, cell_data in enumerate(row_data):
            row_cells[i].text = cell_data
    
    doc.add_paragraph()
    
    # 文本段落测试
    doc.add_heading('4. 文本字段测试', level=1)
    doc.add_paragraph('以下段落包含需要提取的字段信息：')
    doc.add_paragraph('系统配置参数：')
    doc.add_paragraph('• 最大连接数：1000个')
    doc.add_paragraph('• 超时时间：30秒')
    doc.add_paragraph('• 缓冲区大小：8192字节')
    doc.add_paragraph('• 重试次数：3次')
    
    doc.add_paragraph('性能指标：')
    doc.add_paragraph('• 吞吐量：500 Mbps')
    doc.add_paragraph('• 延迟：< 10 ms')
    doc.add_paragraph('• 可用性：99.9%')
    
    # 保存文档
    filename = '字段配置测试文档.docx'
    doc.save(filename)
    print(f"✅ 测试文档已生成: {filename}")
    
    return filename

def create_test_config_files():
    """创建测试用的配置文件"""
    
    # 测试配置1：最小配置
    minimal_config = {
        "protocolFields": [
            {"id": "min_pf_1", "name": "参数"},
            {"id": "min_pf_2", "name": "数据类型"}
        ],
        "targetFields": [
            {"id": "min_tf_1", "name": "字段名"},
            {"id": "min_tf_2", "name": "类型"}
        ]
    }
    
    # 测试配置2：标准配置
    standard_config = {
        "protocolFields": [
            {"id": "std_pf_1", "name": "序号"},
            {"id": "std_pf_2", "name": "参数名称"},
            {"id": "std_pf_3", "name": "数据类型说明"},
            {"id": "std_pf_4", "name": "单位"},
            {"id": "std_pf_5", "name": "备注"}
        ],
        "targetFields": [
            {"id": "std_tf_1", "name": "ID"},
            {"id": "std_tf_2", "name": "参数"},
            {"id": "std_tf_3", "name": "数据类型"},
            {"id": "std_tf_4", "name": "单位"},
            {"id": "std_tf_5", "name": "备注"}
        ]
    }
    
    # 测试配置3：复杂配置
    complex_config = {
        "protocolFields": [
            {"id": "comp_pf_1", "name": "序号"},
            {"id": "comp_pf_2", "name": "参数名称"},
            {"id": "comp_pf_3", "name": "数据类型说明"},
            {"id": "comp_pf_4", "name": "单位"},
            {"id": "comp_pf_5", "name": "备注信息"},
            {"id": "comp_pf_6", "name": "范围"},
            {"id": "comp_pf_7", "name": "精度"},
            {"id": "comp_pf_8", "name": "更新频率"}
        ],
        "targetFields": [
            {"id": "comp_tf_1", "name": "ID"},
            {"id": "comp_tf_2", "name": "参数"},
            {"id": "comp_tf_3", "name": "类型"},
            {"id": "comp_tf_4", "name": "单位"},
            {"id": "comp_tf_5", "name": "备注"},
            {"id": "comp_tf_6", "name": "取值范围"},
            {"id": "comp_tf_7", "name": "精度要求"},
            {"id": "comp_tf_8", "name": "采样频率"}
        ]
    }
    
    # 保存配置文件
    configs = [
        ("test_minimal_config.json", minimal_config),
        ("test_standard_config.json", standard_config),
        ("test_complex_config.json", complex_config)
    ]
    
    for filename, config in configs:
        with open(filename, 'w', encoding='utf-8') as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        print(f"✅ 测试配置文件已生成: {filename}")

def create_test_script():
    """创建测试执行脚本"""
    test_script = '''#!/bin/bash
# 字段配置功能测试执行脚本

echo "🚀 开始字段配置功能测试"

# 检查服务状态
echo "1. 检查服务状态..."
if curl -s http://localhost:5001/health > /dev/null; then
    echo "✅ 后端服务正常"
else
    echo "❌ 后端服务未启动"
    exit 1
fi

if curl -s http://localhost:5174 > /dev/null; then
    echo "✅ 前端服务正常"
else
    echo "❌ 前端服务未启动"
    exit 1
fi

# 运行Python测试
echo "2. 运行功能测试..."
python test_field_config_functionality.py

# 检查测试结果
if [ -f "field_config_test_report.json" ]; then
    echo "3. 测试完成，查看详细报告..."
    cat field_config_test_report.json | jq '.summary'
else
    echo "❌ 测试报告未生成"
    exit 1
fi

echo "🎉 测试执行完成！"
'''
    
    with open('run_field_config_tests.sh', 'w') as f:
        f.write(test_script)
    
    # 添加执行权限
    os.chmod('run_field_config_tests.sh', 0o755)
    print("✅ 测试执行脚本已生成: run_field_config_tests.sh")

def main():
    print("🔧 开始生成字段配置测试资源...")
    
    # 生成测试文档
    doc_file = create_test_document()
    
    # 生成测试配置文件
    create_test_config_files()
    
    # 生成测试执行脚本
    create_test_script()
    
    print("\n📋 生成的测试资源:")
    print(f"📄 测试文档: {doc_file}")
    print("⚙️  测试配置文件:")
    print("   - test_minimal_config.json")
    print("   - test_standard_config.json") 
    print("   - test_complex_config.json")
    print("🏃 测试执行脚本: run_field_config_tests.sh")
    print("🐍 功能测试脚本: test_field_config_functionality.py")
    
    print("\n🎯 使用方法:")
    print("1. 确保前后端服务已启动")
    print("2. 运行测试: ./run_field_config_tests.sh")
    print("3. 或直接运行: python test_field_config_functionality.py")
    print("4. 使用测试文档在前端进行手动验证")

if __name__ == "__main__":
    main()