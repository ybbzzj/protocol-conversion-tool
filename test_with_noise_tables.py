# -*- coding: utf-8 -*-
"""
创建包含干扰表的测试文档并对比修复前后效果
重点验证：
1. 目标消息名称兜底提取能力
2. 预览统计修复（遍历所有有效字段表）
"""
import sys
import os
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 设置标准输出编码为UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

sys.path.insert(0, os.path.dirname(__file__))

from backend.services.table_detector import DocumentParser
from backend.routes.mapping import _extract_fields_and_data_from_raw_tables

def create_test_document_with_noise():
    """创建包含干扰表的测试文档"""
    print("=" * 80)
    print("[创建测试文档] 包含干扰表的测试文档")
    print("=" * 80)
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('测试协议文档（包含干扰表）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 干扰表1：测试指令表 =====
    doc.add_paragraph('表1 测试指令表', style='Heading 2')
    table1 = doc.add_table(rows=4, cols=5)
    table1.style = 'Table Grid'
    
    # 表头
    headers1 = ['测试指令', '周期', '执行时机', '错误处理', '备注']
    for idx, header in enumerate(headers1):
        cell = table1.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # 数据行
    data1 = [
        ['指令1', '100ms', '启动时', '重试3次', '系统自检'],
        ['指令2', '200ms', '运行中', '记录日志', '数据采集'],
        ['指令3', '500ms', '停止时', '发送告警', '状态保存']
    ]
    
    for row_idx, row_data in enumerate(data1):
        for col_idx, value in enumerate(row_data):
            table1.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 干扰表2：帧格式说明表 =====
    doc.add_paragraph('表2 帧格式说明', style='Heading 2')
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'
    
    headers2 = ['帧头', '数据域', '帧尾']
    for idx, header in enumerate(headers2):
        cell = table2.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data2 = [
        ['0xAA', '数据内容', '0x55'],
        ['帧头标识', '有效数据', '帧尾标识'],
        ['1字节', 'N字节', '1字节'],
        ['固定值', '可变长度', '固定值']
    ]
    
    for row_idx, row_data in enumerate(data2):
        for col_idx, value in enumerate(row_data):
            table2.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 目标表：XX装置仿真输入数据 =====
    doc.add_paragraph('表3 XX装置仿真输入数据', style='Heading 2')
    table3 = doc.add_table(rows=4, cols=6)
    table3.style = 'Table Grid'
    
    headers3 = ['序号', '参数', '数据类型', '长度', '单位', '备注']
    for idx, header in enumerate(headers3):
        cell = table3.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data3 = [
        ['1', '仿真状态标志位', 'UCHAR', '1', '-', '仿真运行状态标识'],
        ['2', 'XX计时时间', 'UINTEGER-32', '4', 'ms', '计时器计数值'],
        ['3', '测试结果', 'USHORT', '2', '-', '测试结果数据']
    ]
    
    for row_idx, row_data in enumerate(data3):
        for col_idx, value in enumerate(row_data):
            table3.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 干扰表3：周期设置表 =====
    doc.add_paragraph('表4 周期设置表', style='Heading 2')
    table4 = doc.add_table(rows=4, cols=4)
    table4.style = 'Table Grid'
    
    headers4 = ['周期名称', '周期值', '触发条件', '说明']
    for idx, header in enumerate(headers4):
        cell = table4.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data4 = [
        ['采集周期', '100ms', '定时触发', '数据采集周期'],
        ['处理周期', '200ms', '事件触发', '数据处理周期'],
        ['上报周期', '500ms', '定时触发', '数据上报周期']
    ]
    
    for row_idx, row_data in enumerate(data4):
        for col_idx, value in enumerate(row_data):
            table4.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 有效字段表：测试数据表 =====
    doc.add_paragraph('表5 测试数据表', style='Heading 2')
    table5 = doc.add_table(rows=4, cols=6)
    table5.style = 'Table Grid'
    
    headers5 = ['序号', '参数', '数据类型', '长度', '单位', '备注']
    for idx, header in enumerate(headers5):
        cell = table5.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data5 = [
        ['1', '测试参数1', 'UCHAR', '1', 'V', '测试电压'],
        ['2', '测试参数2', 'USHORT', '2', 'A', '测试电流'],
        ['3', '测试参数3', 'UINTEGER-32', '4', 'Hz', '测试频率']
    ]
    
    for row_idx, row_data in enumerate(data5):
        for col_idx, value in enumerate(row_data):
            table5.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 空表（模拟空数据表） =====
    doc.add_paragraph('表6 空数据表', style='Heading 2')
    table6 = doc.add_table(rows=1, cols=6)
    table6.style = 'Table Grid'
    
    headers6 = ['序号', '参数', '数据类型', '长度', '单位', '备注']
    for idx, header in enumerate(headers6):
        cell = table6.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # 保存文档
    output_path = "word/测试文档_包含干扰表.docx"
    doc.save(output_path)
    
    print(f"[SUCCESS] 测试文档已创建: {output_path}")
    print(f"[文档结构]")
    print(f"  表1: 测试指令表（干扰表）")
    print(f"  表2: 帧格式说明（干扰表）")
    print(f"  表3: XX装置仿真输入数据（目标表）")
    print(f"  表4: 周期设置表（干扰表）")
    print(f"  表5: 测试数据表（有效字段表）")
    print(f"  表6: 空数据表（空表）")
    
    return output_path

def test_before_fix(doc_path):
    """测试修复前的效果"""
    print("\n" + "=" * 80)
    print("[修复前测试] 不使用目标名称兜底提取")
    print("=" * 80)
    
    parser = DocumentParser()
    result = parser.parse(doc_path)
    raw_tables = result['tables']
    
    print(f"[识别结果] 总表格数: {len(raw_tables)}")
    
    # 统计表格类型
    field_def_count = 0
    noise_count = 0
    empty_count = 0
    target_found = False
    
    for idx, table in enumerate(raw_tables):
        table_type = table.get('table_type', 'unknown')
        msg_name = table.get('msg_name', '未知')
        data_rows = table.get('data_rows', [])
        
        print(f"  表{idx+1}: {msg_name} (类型: {table_type}, 数据行: {len(data_rows)})")
        
        if table_type == 'field_def':
            field_def_count += 1
            if 'XX装置仿真输入数据' in msg_name:
                target_found = True
        elif table_type == 'skip':
            noise_count += 1
        elif len(data_rows) == 0:
            empty_count += 1
    
    print(f"\n[统计结果]")
    print(f"  有效字段表数: {field_def_count}")
    print(f"  干扰表数: {noise_count}")
    print(f"  空表数: {empty_count}")
    print(f"  目标表查找: {'✅ 成功' if target_found else '❌ 失败'}")
    
    # 测试预览统计（修复前）
    print(f"\n[预览统计测试] 修复前逻辑")
    extracted_fields, table_data = _extract_fields_and_data_from_raw_tables(raw_tables)
    
    print(f"  提取字段数: {len(extracted_fields)}")
    print(f"  预览表格数: {len(table_data)}")
    
    if len(extracted_fields) == 0:
        print(f"  ❌ 预览统计失败：字段数为零")
    else:
        print(f"  ✅ 预览统计正常")
    
    return {
        'total_tables': len(raw_tables),
        'field_def_count': field_def_count,
        'noise_count': noise_count,
        'empty_count': empty_count,
        'target_found': target_found,
        'extracted_fields': len(extracted_fields),
        'preview_tables': len(table_data)
    }

def test_after_fix(doc_path):
    """测试修复后的效果"""
    print("\n" + "=" * 80)
    print("[修复后测试] 使用目标名称兜底提取")
    print("=" * 80)
    
    target_names = ["XX装置仿真输入数据"]
    parser = DocumentParser(target_message_names=target_names)
    result = parser.parse(doc_path)
    raw_tables = result['tables']
    
    print(f"[识别结果] 总表格数: {len(raw_tables)}")
    print(f"[目标名称] {target_names}")
    
    # 统计表格类型
    field_def_count = 0
    noise_count = 0
    empty_count = 0
    target_found = False
    target_table = None
    
    for idx, table in enumerate(raw_tables):
        table_type = table.get('table_type', 'unknown')
        msg_name = table.get('msg_name', '未知')
        data_rows = table.get('data_rows', [])
        
        print(f"  表{idx+1}: {msg_name} (类型: {table_type}, 数据行: {len(data_rows)})")
        
        if table_type == 'field_def':
            field_def_count += 1
            if 'XX装置仿真输入数据' in msg_name:
                target_found = True
                target_table = table
        elif table_type == 'skip':
            noise_count += 1
        elif len(data_rows) == 0:
            empty_count += 1
    
    print(f"\n[统计结果]")
    print(f"  有效字段表数: {field_def_count}")
    print(f"  干扰表数: {noise_count}")
    print(f"  空表数: {empty_count}")
    print(f"  目标表查找: {'✅ 成功' if target_found else '❌ 失败'}")
    
    # 检查目标表的关键字段
    if target_table:
        print(f"\n[目标表验证] XX装置仿真输入数据")
        data_rows = target_table.get('data_rows', [])
        print(f"  数据行数: {len(data_rows)}")
        
        # 查找关键字段
        expected_fields = ["仿真状态标志位", "XX计时时间"]
        found_fields = []
        
        for row in data_rows:
            for key, value in row.items():
                for field in expected_fields:
                    if field in str(value) and field not in found_fields:
                        found_fields.append(field)
                        # 查找数据类型
                        data_type = row.get('数据类型') or row.get('类型') or 'N/A'
                        print(f"  ✅ 找到字段: {field} / {data_type}")
        
        if len(found_fields) >= 2:
            print(f"  ✅ 关键字段提取完整: {len(found_fields)}/2")
        else:
            print(f"  ⚠️  关键字段提取不完整: {len(found_fields)}/2")
    
    # 测试预览统计（修复后）
    print(f"\n[预览统计测试] 修复后逻辑")
    extracted_fields, table_data = _extract_fields_and_data_from_raw_tables(raw_tables)
    
    print(f"  提取字段数: {len(extracted_fields)}")
    print(f"  预览表格数: {len(table_data)}")
    
    if len(extracted_fields) == 0:
        print(f"  ❌ 预览统计失败：字段数为零")
    else:
        print(f"  ✅ 预览统计正常")
    
    return {
        'total_tables': len(raw_tables),
        'field_def_count': field_def_count,
        'noise_count': noise_count,
        'empty_count': empty_count,
        'target_found': target_found,
        'target_fields_found': len(found_fields) if target_table else 0,
        'extracted_fields': len(extracted_fields),
        'preview_tables': len(table_data)
    }

def compare_results(before_result, after_result):
    """对比修复前后的结果"""
    print("\n" + "=" * 80)
    print("[对比分析] 修复前后效果对比")
    print("=" * 80)
    
    print(f"\n[表格识别对比]")
    print(f"  总表格数: {before_result['total_tables']} → {after_result['total_tables']}")
    print(f"  有效字段表数: {before_result['field_def_count']} → {after_result['field_def_count']}")
    print(f"  干扰表数: {before_result['noise_count']} → {after_result['noise_count']}")
    print(f"  空表数: {before_result['empty_count']} → {after_result['empty_count']}")
    
    print(f"\n[目标表查找对比]")
    before_status = "✅ 成功" if before_result['target_found'] else "❌ 失败"
    after_status = "✅ 成功" if after_result['target_found'] else "❌ 失败"
    print(f"  修复前: {before_status}")
    print(f"  修复后: {after_status}")
    
    if not before_result['target_found'] and after_result['target_found']:
        print(f"  ✅ 兜底提取能力生效：修复前未找到目标表，修复后成功找到")
    elif before_result['target_found'] and after_result['target_found']:
        print(f"  ℹ️  兜底机制未触发：前后都找到目标表")
    else:
        print(f"  ❌ 兜底提取失败：前后都未找到目标表")
    
    print(f"\n[预览统计对比]")
    print(f"  提取字段数: {before_result['extracted_fields']} → {after_result['extracted_fields']}")
    print(f"  预览表格数: {before_result['preview_tables']} → {after_result['preview_tables']}")
    
    if before_result['extracted_fields'] == 0 and after_result['extracted_fields'] > 0:
        print(f"  ✅ 预览统计修复生效：修复前字段数为零，修复后成功提取")
    elif before_result['extracted_fields'] > 0 and after_result['extracted_fields'] > 0:
        print(f"  ℹ️  预览统计正常：前后都能提取字段")
    else:
        print(f"  ❌ 预览统计失败：前后字段数都为零")
    
    # 计算提升幅度
    print(f"\n[提升幅度]")
    
    if before_result['extracted_fields'] > 0:
        field_improvement = (after_result['extracted_fields'] - before_result['extracted_fields']) / before_result['extracted_fields'] * 100
        print(f"  字段数提升: {field_improvement:.1f}%")
    
    if before_result['preview_tables'] > 0:
        preview_improvement = (after_result['preview_tables'] - before_result['preview_tables']) / before_result['preview_tables'] * 100
        print(f"  预览表格数提升: {preview_improvement:.1f}%")
    
    # 总体评估
    print(f"\n[总体评估]")
    
    success_count = 0
    total_tests = 3
    
    # 测试1：目标表查找
    if after_result['target_found']:
        success_count += 1
        print(f"  ✅ 目标表查找成功")
    else:
        print(f"  ❌ 目标表查找失败")
    
    # 测试2：关键字段提取
    if after_result.get('target_fields_found', 0) >= 2:
        success_count += 1
        print(f"  ✅ 关键字段提取完整")
    else:
        print(f"  ❌ 关键字段提取不完整")
    
    # 测试3：预览统计
    if after_result['extracted_fields'] > 0:
        success_count += 1
        print(f"  ✅ 预览统计正常")
    else:
        print(f"  ❌ 预览统计失败")
    
    print(f"\n[测试通过率] {success_count}/{total_tests} = {success_count/total_tests*100:.1f}%")
    
    return success_count == total_tests

def main():
    """主函数"""
    print("=" * 80)
    print("[完整测试] 干扰表场景下的修复前后对比")
    print("=" * 80)
    
    # 创建测试文档
    doc_path = create_test_document_with_noise()
    
    if not os.path.exists(doc_path):
        print(f"[ERROR] 测试文档创建失败")
        return False
    
    # 测试修复前
    before_result = test_before_fix(doc_path)
    
    # 测试修复后
    after_result = test_after_fix(doc_path)
    
    # 对比结果
    all_passed = compare_results(before_result, after_result)
    
    # 生成汇总报告
    print("\n" + "=" * 80)
    print("[汇总报告]")
    print("=" * 80)
    
    if all_passed:
        print(f"[SUCCESS] 所有测试通过，修复效果显著")
        print(f"  ✅ 目标消息名称兜底提取能力正常工作")
        print(f"  ✅ 预览统计修复成功，避免字段数为零")
        print(f"  ✅ 关键字段提取完整，数据完整性良好")
    else:
        print(f"[WARNING] 部分测试未通过，需要进一步优化")
    
    print(f"\n[建议]")
    print(f"  1. 查看生成的测试文档: {doc_path}")
    print(f"  2. 对比修复前后的识别结果")
    print(f"  3. 验证兜底提取能力和预览统计修复")
    
    return all_passed

if __name__ == "__main__":
    try:
        success = main()
        sys.exit(0 if success else 1)
    except Exception as e:
        import traceback
        print(f"\n[ERROR] 测试失败: {e}")
        print(traceback.format_exc())
        sys.exit(1)