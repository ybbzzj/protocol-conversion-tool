# -*- coding: utf-8 -*-
"""
生成"不常见表头"测试文档 + 验证能否通过配置兜底识别。
表头全部为非关键词列（编号/项目/规格/说明/备注/代号/指标/区间），
只能走【配置兜底】通道，用于验证 60% 阈值。
"""
import os
import sys
import logging

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.chdir(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from backend.services.table_detector import DocumentParser

logging.basicConfig(level=logging.WARNING, format='[%(name)s] %(message)s')

OUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'test_docs')
os.makedirs(OUT_DIR, exist_ok=True)


def make_bc(name, header_cols, data_rows=2):
    doc = Document()
    doc.add_paragraph(f'测试表 {name}')
    t = doc.add_table(rows=data_rows + 1, cols=len(header_cols))
    for j, h in enumerate(header_cols):
        t.rows[0].cells[j].text = h
    for r in range(1, data_rows + 1):
        for j in range(len(header_cols)):
            t.rows[r].cells[j].text = f'值{r}-{j}'
    path = os.path.join(OUT_DIR, f'uncommon_{name}.docx')
    doc.save(path)
    return path


def make_a(name, header_cols, data_rows=2):
    """多行表头：行0/1 元数据，行2 真实表头。"""
    doc = Document()
    doc.add_paragraph(f'测试表 {name}')
    n = 2 + 1 + data_rows
    t = doc.add_table(rows=n, cols=len(header_cols))
    meta0 = ['通信帧名称', '某信息', '信息标识', '0x1A2B'] + [''] * (len(header_cols) - 4)
    meta1 = ['信息流向', 'A→B', '发送周期', '100ms'] + [''] * (len(header_cols) - 4)
    for j, v in enumerate(meta0):
        t.rows[0].cells[j].text = v
    for j, v in enumerate(meta1):
        t.rows[1].cells[j].text = v
    for j, h in enumerate(header_cols):
        t.rows[2].cells[j].text = h
    for r in range(3, n):
        for j in range(len(header_cols)):
            t.rows[r].cells[j].text = f'值{r}-{j}'
    path = os.path.join(OUT_DIR, f'uncommon_{name}.docx')
    doc.save(path)
    return path


def verify(path, config_fields, title):
    cfg = [{'table_type': 'field_def', 'required_fields': list(config_fields)}]
    p = DocumentParser(config=cfg)
    res = p.parse(path)
    ok = any(t.get('table_type') == 'field_def' for t in res['tables'])
    print(f'  {"✅" if ok else "❌"} {title}')
    return ok


def main():
    print(f'生成测试文档到: {OUT_DIR}\n')
    cases = [
        ('bc100', make_bc('bc100', ['编号', '项目', '规格', '说明']),
         ['编号', '项目', '规格', '说明'], 'B/C 全部命中(100%)'),
        ('bc75', make_bc('bc75', ['编号', '项目', '规格', '说明']),
         ['编号', '项目', '规格'], 'B/C 4列命中3(75%)'),
        ('bc60', make_bc('bc60', ['编号', '项目', '规格', '说明', '备注']),
         ['编号', '项目', '规格'], 'B/C 5列命中3(60%)'),
        ('bc50', make_bc('bc50', ['编号', '项目', '规格', '说明']),
         ['编号', '项目'], 'B/C 4列命中2(50%)'),
        ('a60', make_a('a60', ['代号', '指标', '区间', '说明', '备注']),
         ['代号', '指标', '区间'], 'A型多行表头 5列命中3(60%)'),
        ('a100', make_a('a100', ['代号', '指标', '区间', '说明', '备注']),
         ['代号', '指标', '区间', '说明', '备注'], 'A型多行表头 全命中(100%)'),
    ]
    all_ok = True
    for name, path, cfg, desc in cases:
        print(f'文档: {os.path.basename(path)}  ({desc})')
        ok = verify(path, cfg, desc)
        all_ok = all_ok and ok
    print('\n' + ('全部通过 ✅' if all_ok else '存在未通过 ❌'))


if __name__ == '__main__':
    main()
