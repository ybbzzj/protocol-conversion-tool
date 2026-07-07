# -*- coding: utf-8 -*-
"""
测试「纯自定义非ID表 + 用户配齐所有字段 → 强制识别」。

构造两张完全不含标准关键词（内容/参数/类型/字节/UINT/FLOAT…）的表：
  - 表1：类型B/C（行0 即表头）
  - 表2：类型A（行0=通信帧名称元数据，表头在行3）
并验证：无配置时被过滤，配齐所有字段后两张都被强制识别为 field_def。
"""
import sys
import os
import logging

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.chdir(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from backend.services.table_detector import DocumentParser

logging.basicConfig(level=logging.INFO, format='[%(name)s] %(message)s')

SYN_DOC = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'synthetic_custom.docx')
CUSTOM_FIELDS = ['编号', '项目', '规格', '说明']


def build_synthetic_doc():
    doc = Document()
    doc.add_paragraph('表1 设备参数定义')  # 前置段落（B/C 表消息名来源）
    t1 = doc.add_table(rows=3, cols=4)
    t1.rows[0].cells[0].text = '编号'
    t1.rows[0].cells[1].text = '项目'
    t1.rows[0].cells[2].text = '规格'
    t1.rows[0].cells[3].text = '说明'
    t1.rows[1].cells[0].text = '1'
    t1.rows[1].cells[1].text = '速度'
    t1.rows[1].cells[2].text = '200'
    t1.rows[1].cells[3].text = '主参数'
    t1.rows[2].cells[0].text = '2'
    t1.rows[2].cells[1].text = '温度'
    t1.rows[2].cells[2].text = '300'
    t1.rows[2].cells[3].text = '次参数'

    doc.add_paragraph('')  # 分隔
    doc.add_paragraph('表2 状态信息定义')  # 前置段落（A 表消息名来源）
    # 类型A：行0 元数据，行1 元数据，行2=真正表头，行3~4 数据
    t2 = doc.add_table(rows=5, cols=4)
    t2.rows[0].cells[0].text = '通信帧名称'
    t2.rows[0].cells[1].text = '某状态信息'
    t2.rows[0].cells[2].text = '信息标识'
    t2.rows[0].cells[3].text = '0x1A2B'
    t2.rows[1].cells[0].text = '信息流向'
    t2.rows[1].cells[1].text = 'A→B'
    t2.rows[1].cells[2].text = '发送周期'
    t2.rows[1].cells[3].text = '100ms'
    t2.rows[2].cells[0].text = '编号'
    t2.rows[2].cells[1].text = '项目'
    t2.rows[2].cells[2].text = '规格'
    t2.rows[2].cells[3].text = '说明'
    t2.rows[3].cells[0].text = '1'
    t2.rows[3].cells[1].text = '速度'
    t2.rows[3].cells[2].text = '200'
    t2.rows[3].cells[3].text = '主参数'
    t2.rows[4].cells[0].text = '2'
    t2.rows[4].cells[1].text = '温度'
    t2.rows[4].cells[2].text = '300'
    t2.rows[4].cells[3].text = '次参数'

    doc.save(SYN_DOC)


def summarize(result, tag):
    tables = result['tables']
    print(f'\n=== {tag} ===')
    print(f'总表格数: {len(tables)}')
    for i, t in enumerate(tables):
        ttype = t.get('table_type', '?')
        if ttype == 'skip':
            print(f'  [{i}] SKIP   para={t.get("preceding_para","")[:30]}')
            continue
        hdr = '|'.join(t.get('headers', [])[:8])
        print(f'  [{i}] {ttype:<10} msg="{t.get("msg_name","")}" headers=[{hdr}] rows={len(t.get("data_rows",[]))}')


def main():
    build_synthetic_doc()

    # 场景A：无配置 → 预期两张都被过滤（skip）
    print('#' * 70)
    print('场景A: 不传任何配置（基线，预期全部 skip）')
    print('#' * 70)
    p0 = DocumentParser(config=None)
    r0 = p0.parse(SYN_DOC)
    summarize(r0, '无配置')

    # 场景B：配齐所有字段 → 预期两张都强制识别为 field_def
    print('\n' + '#' * 70)
    print('场景B: required_fields=全部4字段（预期两张都 field_def）')
    print('#' * 70)
    cfg = [{'table_type': 'field_def', 'required_fields': list(CUSTOM_FIELDS)}]
    p1 = DocumentParser(config=cfg)
    r1 = p1.parse(SYN_DOC)
    summarize(r1, '配齐字段')

    # 断言
    # 说明：被 skip 的表不会进入 result['tables']，负向基线改读权威决策日志 log_records
    field_defs = [t for t in r1['tables'] if t.get('table_type') == 'field_def']
    skips_a = [rec for rec in p0.detector.log_records if rec.get('table_type') == 'skip']
    print('\n' + '=' * 70)
    print('断言结果:')
    print(f'  无配置时被过滤的表数(来自决策日志) = {len(skips_a)} (期望 2)')
    print(f'  配齐字段后识别为 field_def 的表数 = {len(field_defs)} (期望 2)')
    ok = len(skips_a) == 2 and len(field_defs) == 2
    print('  =>', 'PASS ✅' if ok else 'FAIL ❌')
    return ok


if __name__ == '__main__':
    ok = main()
    sys.exit(0 if ok else 1)
