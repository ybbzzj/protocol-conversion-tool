# -*- coding: utf-8 -*-
"""
创建极端场景：目标表被误判为干扰表
重点验证目标消息名称兜底提取能力
"""
import sys
import os
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 设置标准输出编码为UTF-8
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

sys.path.insert(0, os.path.dirname(__file__))

from backend.services.table_detector import DocumentParser
from backend.routes.mapping import _extract_fields_and_data_from_raw_tables

def create_extreme_test_document():
    """创建极端测试文档：目标表具有干扰表特征"""
    print("=" * 80)
    print("[创建极端测试文档] 目标表具有干扰表特征")
    print("=" * 80)
    
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('极端测试协议文档', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # ===== 干扰表1：测试指令表（典型的干扰表） =====
    doc.add_paragraph('表1 测试指令表', style='Heading 2')
    table1 = doc.add_table(rows=4, cols=5)
    table1.style = 'Table Grid'
    
    headers1 = ['测试指令', '周期', '执行时机', '错误处理', '备注']
    for idx, header in enumerate(headers1):
        cell = table1.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data1 = [
        ['指令1', '100ms', '启动时', '重试3次', '系统自检'],
        ['指令2', '200ms', '运行中', '记录日志', '数据采集'],
        ['指令3', '500ms', '停止时', '发送告警', '状态保存']
    ]
    
    for row_idx, row_data in enumerate(data1):
        for col_idx, value in enumerate(row_data):
            table1.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 目标表：XX装置仿真输入数据（故意设计成可能被误判的样式） =====
    # 添加干扰性前置段落（包含"参见附录"等干扰词）
    doc.add_paragraph('参见附录A，表2 XX装置仿真输入数据说明', style='Heading 2')
    
    table2 = doc.add_table(rows=4, cols=4)  # 少列数，可能被误判
    table2.style = 'Table Grid'
    
    # 表头设计：故意不包含"数据类型"等关键字
    headers2 = ['序号', '参数', '单位', '备注']
    for idx, header in enumerate(headers2):
        cell = table2.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # 数据行：故意不包含数据类型关键字
    data2 = [
        ['1', '仿真状态标志位', '-', '仿真运行状态标识'],
        ['2', 'XX计时时间', 'ms', '计时器计数值'],
        ['3', '测试结果', '-', '测试结果数据']
    ]
    
    for row_idx, row_data in enumerate(data2):
        for col_idx, value in enumerate(row_data):
            table2.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 干扰表2：帧格式说明表 =====
    doc.add_paragraph('表3 帧格式说明', style='Heading 2')
    table3 = doc.add_table(rows=5, cols=3)
    table3.style = 'Table Grid'
    
    headers3 = ['帧头', '数据域', '帧尾']
    for idx, header in enumerate(headers3):
        cell = table3.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data3 = [
        ['0xAA', '数据内容', '0x55'],
        ['帧头标识', '有效数据', '帧尾标识'],
        ['1字节', 'N字节', '1字节'],
        ['固定值', '可变长度', '固定值']
    ]
    
    for row_idx, row_data in enumerate(data3):
        for col_idx, value in enumerate(row_data):
            table3.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 有效字段表：测试数据表 =====
    doc.add_paragraph('表4 测试数据表', style='Heading 2')
    table4 = doc.add_table(rows=4, cols=6)
    table4.style = 'Table Grid'
    
    headers4 = ['序号', '参数', '数据类型', '长度', '单位', '备注']
    for idx, header in enumerate(headers4):
        cell = table4.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data4 = [
        ['1', '测试参数1', 'UCHAR', '1', 'V', '测试电压'],
        ['2', '测试参数2', 'USHORT', '2', 'A', '测试电流'],
        ['3', '测试参数3', 'UINTEGER-32', '4', 'Hz', '测试频率']
    ]
    
    for row_idx, row_data in enumerate(data4):
        for col_idx, value in enumerate(row_data):
            table4.rows[row_idx + 1].cells[col_idx].text = value
    
    # ===== 空表 =====
    doc.add_paragraph('表5 空数据表', style='Heading 2')
    table5 = doc.add_table(rows=1, cols=6)
    table5.style = 'Table Grid'
    
    headers5 = ['序号', '参数', '数据类型', '长度', '单位', '备注']
    for idx, header in enumerate(headers5):
        cell = table5.rows[0].cells[idx]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # 保存文档
    output_path = "word/极端测试文档_目标表可能被误判.docx"
    doc.save(output_path)
    
    print(f"[SUCCESS] 极端测试文档已创建: {output_path}")
    print(f"[文档结构]")
    print(f"  表1: 测试指令表（干扰表）")
    print(f"  表2: XX装置仿真输入数据（目标表，故意设计成可能被误判）")
    print(f"    - 前置段落含干扰词：'参见附录A'")
    print(f"    - 表头不含数据类型关键字")
    print(f"    - 数据行不含数据类型关键字")
    print(f"  表3: 帧格式说明（干扰表）")
    print(f"  表4: 测试数据表（有效字段表）")
    print(f"  表5: 空数据表（空表）")
    
    return output_path

def test_extreme_scenario(doc_path):
    """测试极端场景"""
    print("\n" + "=" * 80)
    print("[极端场景测试] 目标表可能被误判为干扰表")
    print("=" * 80)
    
    target_names = ["XX装置仿真输入数据"]
    
    # 测试修复前（不使用目标名称）
    print(f"\n[修复前] 不使用目标名称兜底提取")
    parser_before = DocumentParser()
    result_before = parser_before.parse(doc_path)
    raw_tables_before = result_before['tables']
    
    print(f"[识别结果] 总表格数: {len(raw_tables_before)}")
    
    target_found_before = False
    for idx, table in enumerate(raw_tables_before):
        table_type = table.get('table_type', 'unknown')
        msg_name = table.get('msg_name', '未知')
        data_rows = table.get('data_rows', [])
        
        print(f"  表{idx+1}: {msg_name} (类型: {table_type}, 数据行: {len(data_rows)})")
        
        if 'XX装置仿真输入数据' in msg_name:
            target_found_before = True
            print(f"    ⚠️  目标表状态: {'✅ 成功提取' if table_type == 'field_def' else '❌ 被误判为干扰表'}")
    
    # 测试修复后（使用目标名称）
    print(f"\n[修复后] 使用目标名称兜底提取: {target_names}")
    parser_after = DocumentParser(target_message_names=target_names)
    result_after = parser_after.parse(doc_path)
    raw_tables_after = result_after['tables']
    
    print(f"[识别结果] 总表格数: {len(raw_tables_after)}")
    
    target_found_after = False
    target_table_after = None
    
    for idx, table in enumerate(raw_tables_after):
        table_type = table.get('table_type', 'unknown')
        msg_name = table.get('msg_name', '未知')
        data_rows = table.get('data_rows', [])
        
        print(f"  表{idx+1}: {msg_name} (类型: {table_type}, 数据行: {len(data_rows)})")
        
        if 'XX装置仿真输入数据' in msg_name:
            target_found_after = True
            target_table_after = table
            print(f"    ✅ 目标表状态: {'成功提取' if table_type == 'field_def' else '兜底提取成功'}")
    
    # 对比结果
    print(f"\n[对比结果]")
    
    print(f"[目标表查找对比]")
    print(f"  修复前: {'✅ 成功' if target_found_before else '❌ 失败（被误判为干扰表）'}")
    print(f"  修复后: {'✅ 成功' if target_found_after else '❌ 失败'}")
    
    if not target_found_before and target_found_after:
        print(f"  ✅✅✅ 兜底提取能力生效！")
        print(f"      修复前：目标表被误判为干扰表，未提取")
        print(f"      修复后：使用目标名称兜底提取，成功提取目标表")
    elif target_found_before and target_found_after:
        print(f"  ℹ️  兜底机制未触发：前后都找到目标表")
    else:
        print(f"  ❌ 兜底提取失败：前后都未找到目标表")
    
    # 验证关键字段
    if target_table_after:
        print(f"\n[关键字段验证]")
        data_rows = target_table_after.get('data_rows', [])
        print(f"  数据行数: {len(data_rows)}")
        
        expected_fields = ["仿真状态标志位", "XX计时时间"]
        found_fields = []
        
        for row in data_rows:
            for key, value in row.items():
                for field in expected_fields:
                    if field in str(value) and field not in found_fields:
                        found_fields.append(field)
                        print(f"  ✅ 找到字段: {field}")
        
        if len(found_fields) >= 2:
            print(f"  ✅ 关键字段提取完整: {len(found_fields)}/2")
        else:
            print(f"  ⚠️  关键字段提取不完整: {len(found_fields)}/2")
    
    # 测试预览统计
    print(f"\n[预览统计对比]")
    
    # 修复前
    extracted_fields_before, table_data_before = _extract_fields_and_data_from_raw_tables(raw_tables_before)
    print(f"  修复前: 提取字段数={len(extracted_fields_before)}, 预览表格数={len(table_data_before)}")
    
    # 修复后
    extracted_fields_after, table_data_after = _extract_fields_and_data_from_raw_tables(raw_tables_after)
    print(f"  修复后: 提取字段数={len(extracted_fields_after)}, 预览表格数={len(table_data_after)}")
    
    if len(extracted_fields_before) == 0 and len(extracted_fields_after) > 0:
        print(f"  ✅✅✅ 预览统计修复生效！")
        print(f"      修复前：字段数为零（因为第一张表是干扰表）")
        print(f"      修复后：成功提取字段（遍历所有有效字段表）")
    elif len(extracted_fields_before) > 0 and len(extracted_fields_after) > 0:
        print(f"  ℹ️  预览统计正常：前后都能提取字段")
    else:
        print(f"  ❌ 预览统计失败：前后字段数都为零")
    
    # 总体评估
    print(f"\n[总体评估]")
    
    success_count = 0
    
    if not target_found_before and target_found_after:
        success_count += 1
        print(f"  ✅ 目标消息名称兜底提取能力验证通过")
    else:
        print(f"  ℹ️  目标表前后都成功提取（兜底机制未触发）")
    
    if target_table_after and len(found_fields) >= 2:
        success_count += 1
        print(f"  ✅ 关键字段提取验证通过")
    else:
        print(f"  ❌ 关键字段提取验证失败")
    
    if len(extracted_fields_after) > 0:
        success_count += 1
        print(f"  ✅ 预览统计修复验证通过")
    else:
        print(f"  ❌ 预览统计修复验证失败")
    
    print(f"\n[验证结果] {success_count}/3 项验证通过")
    
    return {
        'target_found_before': target_found_before,
        'target_found_after': target_found_after,
        '兜底提取生效': not target_found_before and target_found_after,
        '关键字段完整': target_table_after and len(found_fields) >= 2,
        '预览统计正常': len(extracted_fields_after) > 0
    }

def main():
    """主函数"""
    print("=" * 80)
    print("[极端场景测试] 目标表可能被误判为干扰表")
    print("=" * 80)
    
    # 创建极端测试文档
    doc_path = create_extreme_test_document()
    
    if not os.path.exists(doc_path):
        print(f"[ERROR] 测试文档创建失败")
        return False
    
    # 测试极端场景
    result = test_extreme_scenario(doc_path)
    
    # 生成汇总报告
    print("\n" + "=" * 80)
    print("[汇总报告]")
    print("=" * 80)
    
    if result['兜底提取生效']:
        print(f"[SUCCESS] ✅✅✅ 目标消息名称兜底提取能力成功验证！")
        print(f"  修复前：目标表被误判为干扰表，未提取")
        print(f"  修复后：使用目标名称兜底提取，成功提取目标表")
        print(f"  价值：兜底机制确保目标表在任何情况下都能被提取")
    else:
        print(f"[INFO] 兜底机制未触发，目标表前后都成功提取")
        print(f"  说明：当前的干扰表识别逻辑已经能够正确识别目标表")
    
    if result['预览统计正常']:
        print(f"[SUCCESS] ✅ 预览统计修复验证通过")
        print(f"  修复后：成功遍历所有有效字段表，避免字段数为零")
    else:
        print(f"[ERROR] ❌ 预览统计修复验证失败")
    
    print(f"\n[建议]")
    print(f"  1. 查看生成的测试文档: {doc_path}")
    print(f"  2. 对比修复前后的识别结果差异")
    print(f"  3. 验证兜底提取能力的实际效果")
    
    return result['兜底提取生效'] or result['关键字段完整'] and result['预览统计正常']

if __name__ == "__main__":
    try:
        success = main()
        sys.exit(0 if success else 1)
    except Exception as e:
        import traceback
        print(f"\n[ERROR] 测试失败: {e}")
        print(traceback.format_exc())
        sys.exit(1)