#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成覆盖各种场景的测试 DOCX 文档。
涵盖：普通字段表、混合表（多种变体）、ID表、端口分配表、bit位定义表、噪声表、边界case。
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), 'test_docs_v2')
os.makedirs(OUTPUT_DIR, exist_ok=True)


def add_table(doc, rows_data, col_widths=None):
    """添加一个表格到文档。rows_data 是二维列表。"""
    n_rows = len(rows_data)
    n_cols = max(len(r) for r in rows_data) if rows_data else 0
    if n_rows == 0 or n_cols == 0:
        return

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(rows_data):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < n_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = str(cell_text) if cell_text is not None else ''

    return table


def add_para(doc, text):
    """添加段落"""
    p = doc.add_paragraph(text)
    return p


def make_doc(filename):
    """创建文档并返回"""
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(10.5)
    return doc


def save_doc(doc, filename):
    """保存文档"""
    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    print(f"  生成: {filename}")
    return path


# ─── 1. 普通字段定义表 ──────────────────────────────────────────────────────

def gen_01_simple_standard():
    """普通字段表 - 标准列名（序号/内容/类型/值域/单位/说明）"""
    doc = make_doc('01_simple_standard')
    add_para(doc, '表1 XX装置状态数据')
    add_table(doc, [
        ['序号', '内容', '类型', '值域', '单位', '说明'],
        ['1', 'XX计时时间', 'UINTEGER-32', '0~4294967295', 'ms', '32位整型数'],
        ['2', '测试结果', 'USHORT', '0~65535', '-', '无符号短整型'],
        ['3', '自检状态字1', 'USHORT', '-', '-', '见表A.10'],
        ['4', 'CRC校验码', 'USHORT', '-', '-', 'CRC校验'],
    ])
    save_doc(doc, '01_simple_standard.docx')


def gen_02_simple_no_seq():
    """普通字段表 - 无序号列（内容/类型/长度）"""
    doc = make_doc('02_simple_no_seq')
    add_para(doc, '表2 控制指令数据')
    add_table(doc, [
        ['内容', '数据类型', '数据长度（字节）', '值域', '说明'],
        ['帧头', '0xEB 0x96', '2', '-', '固定标识'],
        ['帧长', 'USHORT', '2', '-', '总字节数'],
        ['校验和', 'USHORT', '2', '-', 'CRC校验'],
    ])
    save_doc(doc, '02_simple_no_seq.docx')


def gen_03_simple_non_standard():
    """普通字段表 - 非标准列名（需要配置才能识别）"""
    doc = make_doc('03_simple_non_standard')
    add_para(doc, '表3 传感器参数表')
    add_table(doc, [
        ['编号', '项目', '规格', '量程', '单位', '备注'],
        ['1', '温度传感器', 'PT100', '-50~200', '℃', '铂电阻'],
        ['2', '压力传感器', '压阻式', '0~10', 'MPa', ''],
        ['3', '湿度传感器', '电容式', '0~100', '%RH', ''],
    ])
    save_doc(doc, '03_simple_non_standard.docx')


def gen_04_simple_with_info_name_col():
    """普通字段表 - 首列名为"信息名称"（不是元数据行，是列名）"""
    doc = make_doc('04_simple_info_name_col')
    add_para(doc, '表4 指令定义表')
    add_table(doc, [
        ['信息名称', '内容', '数据类型', '长度', '说明'],
        ['启动指令', '设备启动', 'UINT8', '1', '0x01=启动'],
        ['停止指令', '设备停止', 'UINT8', '1', '0x02=停止'],
        ['复位指令', '设备复位', 'UINT8', '1', '0x03=复位'],
    ])
    save_doc(doc, '04_simple_info_name_col.docx')


# ─── 2. 混合表 Type A（信息名称行 + 元数据 + 表头） ──────────────────────────

def gen_05_mixed_standard():
    """混合表 - 标准格式（信息名称/信源信宿/传输周期/发起时机 + 序号/内容/类型）"""
    doc = make_doc('05_mixed_standard')
    add_para(doc, '表5 自检结果信息')
    add_table(doc, [
        ['信息名称', '信息名称', '自检结果', '自检结果', '信息标识'],
        ['信源、信宿', '信源、信宿', 'RT13-SA15-8BC', 'RT13-SA15-8BC', 'RT13-SA15-8BC'],
        ['传输周期', '传输周期', '非周期', '非周期', '非周期'],
        ['发起时机', '发起时机', '上电后自检', '上电后自检', ''],
        ['序号', '内容', '内容', '类型', '类型'],
        ['', 'XX计时时间', 'XX计时时间', 'UINTEGER-32', 'UINTEGER-32'],
        ['', '测试结果', '测试结果', 'USHORT', 'USHORT'],
        ['', '自检状态字1', '自检状态字1', 'USHORT', 'USHORT'],
        ['', 'CRC校验码', 'CRC校验码', 'USHORT', 'USHORT'],
    ])
    save_doc(doc, '05_mixed_standard.docx')


def gen_06_mixed_comm_frame_name():
    """混合表 - 使用"通信帧名字"而非"信息名称"（非默认关键词）"""
    doc = make_doc('06_mixed_comm_frame')
    add_para(doc, '表6 自检结果')
    add_table(doc, [
        ['通信帧名字', '通信帧名字', '自检结果', '自检结果', '信息流向'],
        ['前置条件', '前置条件', '非周期', '非周期', '非周期'],
        ['错误处理', '错误处理', '错误处理', '错误处理', '错误处理'],
        ['序号', '内容', '内容', '长度', '长度'],
        ['', 'XX计时时间', 'XX计时时间', '32', '32'],
        ['', '测试结果', '测试结果', '8', '8'],
        ['', 'CRC校验码', 'CRC校验码', '8', '8'],
    ])
    save_doc(doc, '06_mixed_comm_frame.docx')


def gen_07_mixed_type_b():
    """混合表 Type B - 上级信息名称格式（表头在行1）"""
    doc = make_doc('07_mixed_type_b')
    add_para(doc, '表7 XX装置数据1')
    add_table(doc, [
        ['信息名称', '信息名称', 'XX装置数据1', 'XX装置数据1', '上级信息名称'],
        ['序号', '参数', '参数', '数据类型', '数据长度（字节）'],
        ['', 'XX计时时间', 'XX计时时间', 'UINTEGER-32', '4'],
        ['', '帧计数', '帧计数', 'UINTEGER-32', '4'],
    ])
    save_doc(doc, '07_mixed_type_b.docx')


def gen_08_mixed_non_standard_header():
    """混合表 - 表头使用非标准列名（需要配置）"""
    doc = make_doc('08_mixed_non_std')
    add_para(doc, '表8 传感器校准数据')
    add_table(doc, [
        ['信息名称', '信息名称', '校准参数', '校准参数', '信息标识'],
        ['信源、信宿', '信源、信宿', 'RT5-SA3-12BC', 'RT5-SA3-12BC', 'RT5-SA3-12BC'],
        ['传输周期', '传输周期', '100ms', '100ms', '100ms'],
        ['编号', '项目', '项目', '规格', '规格'],
        ['', '零点校准', '零点校准', '0.000', '0.000'],
        ['', '满量程校准', '满量程校准', '10.000', '10.000'],
        ['', '温度补偿', '温度补偿', '25.0', '25.0'],
    ])
    save_doc(doc, '08_mixed_non_std.docx')


def gen_09_mixed_minimal_data():
    """混合表 - 只有1行数据（最小数据块）"""
    doc = make_doc('09_mixed_minimal')
    add_para(doc, '表9 状态数据')
    add_table(doc, [
        ['信息名称', '信息名称', '状态数据', '状态数据', '信息标识'],
        ['信源、信宿', '信源、信宿', 'RT2-SA1-3BC', 'RT2-SA1-3BC', 'RT2-SA1-3BC'],
        ['传输周期', '传输周期', '50ms', '50ms', '50ms'],
        ['序号', '内容', '内容', '类型', '类型'],
        ['', '设备状态', '设备状态', 'USHORT', 'USHORT'],
    ])
    save_doc(doc, '09_mixed_minimal.docx')


def gen_10_mixed_many_metadata():
    """混合表 - 元数据行很多（5行元数据 + 表头）"""
    doc = make_doc('10_mixed_many_meta')
    add_para(doc, '表10 扩展状态信息')
    add_table(doc, [
        ['信息名称', '信息名称', '扩展状态', '扩展状态', '信息标识'],
        ['信源、信宿', '信源、信宿', 'RT8-SA2-5BC', 'RT8-SA2-5BC', 'RT8-SA2-5BC'],
        ['传输周期', '传输周期', '200ms', '200ms', '200ms'],
        ['发起时机', '发起时机', '上电后', '上电后', '上电后'],
        ['前置条件', '前置条件', '初始化完成', '初始化完成', '初始化完成'],
        ['错误处理', '错误处理', '重发3次', '重发3次', '重发3次'],
        ['备注', '备注', '含校验', '含校验', '含校验'],
        ['序号', '内容', '内容', '类型', '值域'],
        ['', '工作模式', '工作模式', 'USHORT', '0~3'],
        ['', '故障码', '故障码', 'USHORT', '-'],
        ['', 'CRC校验码', 'CRC校验码', 'USHORT', '-'],
    ])
    save_doc(doc, '10_mixed_many_meta.docx')


def gen_11_mixed_field_like_metadata():
    """混合表 - 元数据行包含字段类似的词（发送周期/100ms）"""
    doc = make_doc('11_mixed_field_like')
    add_para(doc, '表11 测量数据')
    add_table(doc, [
        ['信息名称', '信息名称', '测量数据', '测量数据', '测量数据'],
        ['信源、信宿', '信源、信宿', 'RT13-SA1-21BC', 'RT13-SA1-21BC', 'RT13-SA1-21BC'],
        ['传输周期', '传输周期', '5ms', '5ms', '其他'],
        ['发起时机', '发起时机', '上电后', '上电后', '错误处理'],
        ['序号', '内容', '内容', '类型', '类型'],
        ['', 'XX计时时间', 'XX计时时间', 'UINTEGER-32', 'UINTEGER-32'],
        ['', 'CRC校验码', 'CRC校验码', 'USHORT', 'USHORT'],
    ])
    save_doc(doc, '11_mixed_field_like.docx')


def gen_12_mixed_type_value_unit():
    """混合表 - 表头列不同（类型/值域/单位/数据处理方法）"""
    doc = make_doc('12_mixed_type_value')
    add_para(doc, '表12 测试结果数据')
    add_table(doc, [
        ['信息名称', '信息名称', '测试结果', '测试结果', '信息标识'],
        ['信源、信宿', 'RT13-SA16-29BC', 'RT13-SA16-29BC', 'RT13-SA16-29BC', 'RT13-SA16-29BC'],
        ['200ms', '200ms', '其他', '其他', '-'],
        ['', '错误处理', '错误处理', '-', '-'],
        ['类型', '值域', '值域', '单位', '数据处理方法'],
        ['UINTEGER-32', '0~4294967295', '0~4294967295', 'ms', '32位整型数'],
        ['USHORT', '-', '-', '-', '每位0为无故障'],
        ['USHORT', '-', '-', '-', 'CRC校验'],
    ])
    save_doc(doc, '12_mixed_type_value.docx')


# ─── 3. ID表 ──────────────────────────────────────────────────────────────

def gen_13_id_standard():
    """ID表 - 标准格式（ID序号/ID定义/是否有数据）"""
    doc = make_doc('13_id_standard')
    add_para(doc, '表13 消息ID分配表')
    add_table(doc, [
        ['ID序号', 'ID定义', '是否有数据'],
        ['0x0301', 'XX装置仿真输入数据', '是'],
        ['0x0302', 'XX装置仿真输出数据', '是'],
        ['0x0A01', 'XX装置自检结果', '是'],
    ])
    save_doc(doc, '13_id_standard.docx')


def gen_14_id_old_format():
    """ID表 - 旧格式（消息ID/信息内容）"""
    doc = make_doc('14_id_old')
    add_para(doc, '表14 消息ID映射')
    add_table(doc, [
        ['序号', '信源', '信宿', '信息内容', '消息ID'],
        ['', '', '', 'XX装置数据1', '0x8000'],
        ['', '', '', '测试数据2', '0x8001'],
        ['', '', '', '状态数据3', '0x8002'],
    ])
    save_doc(doc, '14_id_old.docx')


def gen_15_id_non_standard():
    """ID表 - 非标准列名（需要配置）"""
    doc = make_doc('15_id_non_std')
    add_para(doc, '表15 指令编号表')
    add_table(doc, [
        ['编号', '名称', '有无数据'],
        ['0x0101', '启动指令', '有'],
        ['0x0102', '停止指令', '有'],
        ['0x0103', '复位指令', '无'],
    ])
    save_doc(doc, '15_id_non_standard.docx')


# ─── 4. 端口分配表 ──────────────────────────────────────────────────────────

def gen_16_port_allocation():
    """端口分配表"""
    doc = make_doc('16_port_alloc')
    add_para(doc, '表16 端口分配表')
    add_table(doc, [
        ['序号', '信源系统码', '信宿系统码', '信息内容', '接收组播地址'],
        ['1', '0x01', '0x02', '状态数据', '239.1.1.1'],
        ['2', '0x03', '0x04', '控制指令', '239.1.1.2'],
    ])
    save_doc(doc, '16_port_allocation.docx')


# ─── 5. Bit位定义表 ─────────────────────────────────────────────────────────

def gen_17_bit_def():
    """Bit位定义表"""
    doc = make_doc('17_bit_def')
    add_para(doc, '表17 自检状态字定义')
    add_table(doc, [
        ['类别', '位号', '参数', '含义'],
        ['检测项目', 'D0', 'CPU', '0：未通过，1：通过'],
        ['检测项目', 'D1', 'Flash', '0：未通过，1：通过'],
        ['检测项目', 'D2', 'UART', '0：未通过，1：通过'],
        ['检测项目', 'D3', 'F', '0：未通过，1：通过'],
        ['检测项目', 'D4', 'MU', '0：未通过，1：通过'],
    ])
    save_doc(doc, '17_bit_def.docx')


# ─── 6. 噪声表 ──────────────────────────────────────────────────────────────

def gen_18_noise_cover():
    """噪声表 - 封面表（全空或测试数据）"""
    doc = make_doc('18_noise_cover')
    add_table(doc, [
        ['', '', '', '', ''],
        ['', '', '', '', '公开'],
        ['', '', '', '', ''],
        ['', '测试数据', '测试数据', '测试数据', '测试数据'],
        ['会签', '测试数据', '测试数据', '测试数据', '测试数据'],
        ['', '', '', '', ''],
    ])
    save_doc(doc, '18_noise_cover.docx')


def gen_19_noise_scenario():
    """噪声表 - 场景验证表"""
    doc = make_doc('19_noise_scenario')
    add_para(doc, '表19 验证场景')
    add_table(doc, [
        ['场景', '充分性要求', '验证内容', '验证形式'],
        ['场景1', '', '', ''],
        ['场景2', '', '', ''],
        ['场景3', '', '', ''],
    ])
    save_doc(doc, '19_noise_scenario.docx')


def gen_20_noise_frame_format():
    """噪声表 - 帧格式说明表"""
    doc = make_doc('20_noise_frame')
    add_para(doc, '表20 帧格式')
    add_table(doc, [
        ['字节号', '字段', '说明'],
        ['0~1', '帧头', '0xEB 0x96'],
        ['2~3', '帧长', '总字节数'],
        ['4~5', '帧尾', '0x0D 0x0A'],
    ])
    save_doc(doc, '20_noise_frame_format.docx')


def gen_21_noise_bit_labels():
    """噪声表 - D0/D1位标签表"""
    doc = make_doc('21_noise_bit_labels')
    add_table(doc, [
        ['D0', 'D1', 'D2', 'D3', 'D4'],
        ['', '', '', '', ''],
        ['D8', 'D9', 'D10', 'D11', 'D12'],
        ['', '', '', '', ''],
    ])
    save_doc(doc, '21_noise_bit_labels.docx')


# ─── 7. 边界 case ───────────────────────────────────────────────────────────

def gen_22_edge_2col_header():
    """边界case - 只有2列的表头"""
    doc = make_doc('22_edge_2col')
    add_para(doc, '表22 字节定义')
    add_table(doc, [
        ['字节号', '字段'],
        ['0', '帧头'],
        ['1', '命令字'],
        ['2', '数据长度'],
        ['3', '校验和'],
    ])
    save_doc(doc, '22_edge_2col_header.docx')


def gen_23_edge_coordinate():
    """边界case - 坐标系表（噪声表，无字段关键词）"""
    doc = make_doc('23_edge_coord')
    add_table(doc, [
        ['坐标系1', '动作及时间', '检查项目', '坐标系2', '状态'],
        ['', '', '', '', ''],
        ['', '', '', '', ''],
    ])
    save_doc(doc, '23_edge_coordinate.docx')


def gen_24_edge_action_table():
    """边界case - 动作表（序号/位置状态/动作，无数据类型关键词）"""
    doc = make_doc('24_edge_action')
    add_para(doc, '表24 动作序列')
    add_table(doc, [
        ['序号', '位置状态', '动作'],
        ['1', '初始位置', '等待'],
        ['2', '位置A', '移动到A'],
        ['3', '位置B', '移动到B'],
    ])
    save_doc(doc, '24_edge_action_table.docx')


def gen_25_mixed_with_notes():
    """混合表 - 末尾有注释行"""
    doc = make_doc('25_mixed_notes')
    add_para(doc, '表25 运行状态数据')
    add_table(doc, [
        ['信息名称', '信息名称', '运行状态数据', '运行状态数据', '信息标识'],
        ['信源、信宿', '信源、信宿', 'RT13-SA5-11BC', 'RT13-SA5-11BC', 'RT13-SA5-11BC'],
        ['传输周期', '传输周期', '200ms', '200ms', '其他'],
        ['发起时机', '发起时机', '', '', '错误处理'],
        ['序号', '内容', '内容', '类型', '值域'],
        ['', 'XX计时时间', 'XX计时时间', 'UINTEGER-32', '0~4294967295'],
        ['', '执行的指令或状态', '执行的指令或状态', 'USHORT', '-'],
        ['', 'CRC校验码', 'CRC校验码', 'USHORT', '-'],
        ['注1：以上电为零点。', '注1：以上电为零点。', '注1：以上电为零点。', '注1：以上电为零点。', '注1：以上电为零点。'],
    ])
    save_doc(doc, '25_mixed_with_notes.docx')


def gen_26_mixed_different_col_count():
    """混合表 - 元数据行和表头行列数不同"""
    doc = make_doc('26_mixed_diff_cols')
    add_para(doc, '表26 命令参数')
    add_table(doc, [
        ['信息名称', '命令参数', '信息标识'],
        ['信源、信宿', 'RT3-SA1-2BC', 'RT3-SA1-2BC'],
        ['传输周期', '非周期', '非周期'],
        ['序号', '内容', '类型', '值域', '单位', '说明'],
        ['', '命令字', 'UINT8', '0~255', '-', '命令编号'],
        ['', '参数1', 'UINT16', '0~65535', '-', '参数值1'],
        ['', '参数2', 'UINT16', '0~65535', '-', '参数值2'],
    ])
    save_doc(doc, '26_mixed_diff_cols.docx')


def gen_27_mixed_with_config_fields():
    """混合表 - 表头全为自定义词，需要配置才能识别"""
    doc = make_doc('27_mixed_config')
    add_para(doc, '表27 自定义协议参数')
    add_table(doc, [
        ['信息名称', '信息名称', '自定义参数', '自定义参数', '信息标识'],
        ['信源、信宿', '信源、信宿', 'RT6-SA4-10BC', 'RT6-SA4-10BC', 'RT6-SA4-10BC'],
        ['传输周期', '传输周期', '50ms', '50ms', '50ms'],
        ['代号', '描述', '描述', '规格', '规格'],
        ['', 'A01', 'A01', '100', '100'],
        ['', 'B02', 'B02', '200', '200'],
        ['', 'C03', 'C03', '300', '300'],
    ])
    save_doc(doc, '27_mixed_with_config_fields.docx')


def gen_28_all_in_one():
    """综合文档 - 包含多种表格类型"""
    doc = make_doc('28_all_in_one')

    # 封面噪声表
    add_table(doc, [
        ['', '', '', '', ''],
        ['', '', '', '', '公开'],
        ['', '测试数据', '测试数据', '测试数据', '测试数据'],
    ])

    # ID表
    add_para(doc, '表A 消息ID分配')
    add_table(doc, [
        ['ID序号', 'ID定义', '是否有数据'],
        ['0x0301', 'XX装置输入数据', '是'],
        ['0x0302', 'XX装置输出数据', '是'],
    ])

    # 混合表
    add_para(doc, '表B 自检结果')
    add_table(doc, [
        ['信息名称', '信息名称', '自检结果', '自检结果', '信息标识'],
        ['信源、信宿', '信源、信宿', 'RT13-SA15-8BC', 'RT13-SA15-8BC', 'RT13-SA15-8BC'],
        ['传输周期', '传输周期', '非周期', '非周期', '非周期'],
        ['序号', '内容', '内容', '类型', '类型'],
        ['', 'XX计时时间', 'XX计时时间', 'UINTEGER-32', 'UINTEGER-32'],
        ['', '测试结果', '测试结果', 'USHORT', 'USHORT'],
    ])

    # 普通字段表
    add_para(doc, '表C 字节定义')
    add_table(doc, [
        ['字节号', '意义', '说明'],
        ['0', '帧头', '0xEB'],
        ['1', '命令', '0x01'],
    ])

    # bit位定义表
    add_para(doc, '表D 状态字定义')
    add_table(doc, [
        ['类别', '位号', '参数', '含义'],
        ['检测', 'D0', 'CPU', '0：失败 1：通过'],
        ['检测', 'D1', 'RAM', '0：失败 1：通过'],
    ])

    # 场景噪声表
    add_para(doc, '表E 验证场景')
    add_table(doc, [
        ['场景', '充分性要求', '验证内容', '验证形式'],
        ['', '', '', ''],
    ])

    save_doc(doc, '28_all_in_one.docx')


def gen_29_mixed_empty_metadata_value():
    """混合表 - 元数据值有时为空"""
    doc = make_doc('29_mixed_empty_meta')
    add_para(doc, '表29 遥测数据')
    add_table(doc, [
        ['信息名称', '信息名称', '遥测数据', '遥测数据', '信息标识'],
        ['信源、信宿', '信源、信宿', 'RT1-SA2-3BC', 'RT1-SA2-3BC', 'RT1-SA2-3BC'],
        ['传输周期', '传输周期', '100ms', '100ms', '100ms'],
        ['发起时机', '发起时机', '', '', ''],
        ['错误处理', '错误处理', '重发', '重发', '重发'],
        ['序号', '内容', '内容', '类型', '类型'],
        ['', '电压', '电压', 'FLOAT', 'FLOAT'],
        ['', '电流', '电流', 'FLOAT', 'FLOAT'],
        ['', '温度', '温度', 'USHORT', 'USHORT'],
    ])
    save_doc(doc, '29_mixed_empty_metadata_value.docx')


def gen_30_simple_data_only():
    """普通表 - 只有数据类型列，无序号无内容"""
    doc = make_doc('30_simple_data_only')
    add_para(doc, '表30 参数定义')
    add_table(doc, [
        ['参数', '数据类型', '数据长度（字节）', '值域', '单位'],
        ['温度', 'FLOAT', '4', '-50~200', '℃'],
        ['湿度', 'FLOAT', '4', '0~100', '%RH'],
        ['气压', 'FLOAT', '4', '800~1100', 'hPa'],
    ])
    save_doc(doc, '30_simple_data_only.docx')


if __name__ == '__main__':
    print("生成测试文档到:", OUTPUT_DIR)
    generators = [
        gen_01_simple_standard,
        gen_02_simple_no_seq,
        gen_03_simple_non_standard,
        gen_04_simple_with_info_name_col,
        gen_05_mixed_standard,
        gen_06_mixed_comm_frame_name,
        gen_07_mixed_type_b,
        gen_08_mixed_non_standard_header,
        gen_09_mixed_minimal_data,
        gen_10_mixed_many_metadata,
        gen_11_mixed_field_like_metadata,
        gen_12_mixed_type_value_unit,
        gen_13_id_standard,
        gen_14_id_old_format,
        gen_15_id_non_standard,
        gen_16_port_allocation,
        gen_17_bit_def,
        gen_18_noise_cover,
        gen_19_noise_scenario,
        gen_20_noise_frame_format,
        gen_21_noise_bit_labels,
        gen_22_edge_2col_header,
        gen_23_edge_coordinate,
        gen_24_edge_action_table,
        gen_25_mixed_with_notes,
        gen_26_mixed_different_col_count,
        gen_27_mixed_with_config_fields,
        gen_28_all_in_one,
        gen_29_mixed_empty_metadata_value,
        gen_30_simple_data_only,
    ]
    for gen in generators:
        gen()
    print(f"\n共生成 {len(generators)} 个测试文档")
