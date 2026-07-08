# -*- coding: utf-8 -*-
"""
探测「自定义非ID表 + 配置兜底」的通过范围。
用完全不常见的表头（无 内容/参数/类型/字节 等标准词），
在多个匹配率下验证：配齐字段后能否被强制识别为 field_def。

匹配率 = 表头中命中配置字段的列数 / 表头总列数（精确匹配口径）。
"""
import sys
import os
import logging
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.chdir(os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from backend.services.table_detector import DocumentParser

logging.basicConfig(level=logging.WARNING, format='[%(name)s] %(message)s')


def build_doc(header_cols, kind, title):
    """kind='bc' 行0即表头；kind='a' 行0/1为元数据、真实表头在下方。"""
    doc = Document()
    doc.add_paragraph(title)
    if kind == 'bc':
        n_rows = 3
        t = doc.add_table(rows=n_rows, cols=len(header_cols))
        for j, h in enumerate(header_cols):
            t.rows[0].cells[j].text = h
        # 两行无意义数据（避免空表被跳过）
        for r in (1, 2):
            for j in range(len(header_cols)):
                t.rows[r].cells[j].text = f'v{r}{j}'
    else:  # 'a' 多行表头
        # 行0 元数据 / 行1 元数据 / 行2 真实表头 / 行3~4 数据
        n_rows = 2 + 1 + 2
        t = doc.add_table(rows=n_rows, cols=len(header_cols))
        meta0 = ['通信帧名称', '某信息', '信息标识', '0x1A2B'] + [''] * (len(header_cols) - 4)
        meta1 = ['信息流向', 'A→B', '发送周期', '100ms'] + [''] * (len(header_cols) - 4)
        for j, v in enumerate(meta0):
            t.rows[0].cells[j].text = v
        for j, v in enumerate(meta1):
            t.rows[1].cells[j].text = v
        for j, h in enumerate(header_cols):
            t.rows[2].cells[j].text = h
        for r in (3, 4):
            for j in range(len(header_cols)):
                t.rows[r].cells[j].text = f'v{r}{j}'
    fd, path = tempfile.mkstemp(suffix='.docx', prefix='fb_range_')
    os.close(fd)
    doc.save(path)
    return path


def decision(header_cols, config_fields, kind, title):
    path = build_doc(header_cols, kind, title)
    cfg = [{'table_type': 'field_def', 'required_fields': list(config_fields)}]
    p = DocumentParser(config=cfg)
    res = p.parse(path)
    os.remove(path)
    # 真实判据：表是否真正被提取进 result['tables']（log_records 仅分类，可能返回 skip 不提取）
    identified = any(t.get('table_type') == 'field_def' for t in res['tables'])
    dec = 'field_def' if identified else 'skip'
    # 计算匹配率
    hset = set(header_cols)
    cset = set(config_fields)
    hit = len(hset & cset)
    rate = hit / max(len(header_cols), 1)
    return dec, hit, len(header_cols), rate


def main():
    # 表头全部为"非关键词"列（不含 内容/参数/名称/类型/字节/序号/单位/说明 等），
    # 确保只能走【配置兜底】通道，从而真实探测兜底阈值。
    scenarios = [
        # (表头, 配置字段, 类型, 标题, 备注)
        (['编号', '项目', '规格', '说明'], ['编号', '项目', '规格', '说明'], 'bc', '表-100%', '全部命中'),
        (['编号', '项目', '规格', '说明'], ['编号', '项目', '规格'], 'bc', '表-75%', '4列命中3'),
        (['编号', '项目', '规格', '说明', '备注'], ['编号', '项目', '规格'], 'bc', '表-60%', '5列命中3'),
        (['编号', '项目', '规格', '说明'], ['编号', '项目'], 'bc', '表-50%', '4列命中2'),
        (['代号', '指标', '区间', '说明', '备注'], ['代号', '指标', '区间'], 'a', 'A表-60%', '多行表头含说明(候选词)'),
        (['代号', '指标', '区间', '说明', '备注'], ['代号', '指标', '区间', '说明', '备注'], 'a', 'A表-100%', '多行表头全命中'),
    ]
    print(f'{"场景":<12} {"表型":<4} {"命中/总":<7} {"匹配率":<7} {"决策":<10} 说明')
    print('-' * 78)
    results = []
    for hdr, cfg, kind, title, note in scenarios:
        dec, hit, total, rate = decision(hdr, cfg, kind, title)
        flag = 'PASS ✅' if dec == 'field_def' else 'FAIL ❌'
        print(f'{title:<12} {kind:<4} {hit}/{total:<5} {rate*100:>5.0f}%  {dec:<10} {note}  {flag}')
        results.append((title, rate, dec))
    return results


if __name__ == '__main__':
    main()
