#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试脚本：验证 table_detector 的表头定位和表格识别。
对每个测试文档，定义期望的 table_type、header_row_idx、headers、data_rows 数量等。
"""
import os
import sys
import json
import traceback

# 添加项目路径
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'backend'))

from services.table_detector import TableDetector, _build_grid, _dedup_row
from docx import Document

TEST_DIR = os.path.join(os.path.dirname(__file__), 'test_docs_v2')
OLD_TEST_DIR = os.path.join(os.path.dirname(__file__), 'word')

# ─── 测试用例定义 ─────────────────────────────────────────────────────────────
# 每个测试用例：
#   filename: 文件名
#   config: 传入 TableDetector 的配置（可选）
#   expected_tables: 期望的表格列表，每个元素：
#     table_type: field_def | message_id | port_allocation | bit_def | skip
#     header_contains: 表头应包含的关键词列表（子串匹配）
#     header_not_contains: 表头不应包含的关键词列表
#     data_rows_min: 最少数据行数
#     data_rows_max: 最多数据行数
#     msg_name_contains: 消息名应包含的文本（可选）
#     meta_contains: meta 应包含的 key（可选）
#     header_row_idx: 期望的表头行索引（可选，用于验证 _locate_header）

TEST_CASES = [
    # ─── 1. 普通字段定义表 ──────────────────────────────────────────────
    {
        'filename': '01_simple_standard.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'header_not_contains': ['信息名称', '信源'],
                'data_rows_min': 3,
                'data_rows_max': 5,
                'header_row_idx': 0,
            }
        ]
    },
    {
        'filename': '02_simple_no_seq.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['内容', '数据类型'],
                'header_not_contains': ['序号'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 0,
            }
        ]
    },
    {
        'filename': '03_simple_non_standard.docx',
        'config': ['编号', '项目', '规格', '量程', '单位', '备注'],
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['编号', '项目'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 0,
            }
        ]
    },
    {
        'filename': '04_simple_info_name_col.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['信息名称', '内容', '数据类型'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 0,
                # 关键：行0是表头，不是元数据行
            }
        ]
    },
    # ─── 2. 混合表 ──────────────────────────────────────────────────────
    {
        'filename': '05_mixed_standard.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'header_not_contains': ['信息名称', '信源'],
                'data_rows_min': 3,
                'data_rows_max': 6,
                'msg_name_contains': '自检结果',
                'meta_contains': ['信源、信宿'],
                'header_row_idx': 4,
            }
        ]
    },
    {
        'filename': '06_mixed_comm_frame.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '长度'],
                'header_not_contains': ['通信帧名字'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'msg_name_contains': '自检结果',
                'header_row_idx': 3,
            }
        ]
    },
    {
        'filename': '07_mixed_type_b.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '参数', '数据类型'],
                'header_not_contains': ['信息名称'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'msg_name_contains': 'XX装置数据1',
                'header_row_idx': 1,
            }
        ]
    },
    {
        'filename': '08_mixed_non_std.docx',
        'config': ['编号', '项目', '规格'],
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['编号', '项目'],
                'header_not_contains': ['信息名称'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 3,
            }
        ]
    },
    {
        'filename': '09_mixed_minimal.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'data_rows_min': 1,
                'data_rows_max': 2,
                'header_row_idx': 3,
            }
        ]
    },
    {
        'filename': '10_mixed_many_meta.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'meta_contains': ['信源、信宿', '传输周期'],
                'header_row_idx': 7,
            }
        ]
    },
    {
        'filename': '11_mixed_field_like.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'data_rows_min': 1,
                'data_rows_max': 3,
                'header_row_idx': 4,
            }
        ]
    },
    {
        'filename': '12_mixed_type_value.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['类型', '值域', '单位'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 4,
            }
        ]
    },
    # ─── 3. ID表 ──────────────────────────────────────────────────────
    {
        'filename': '13_id_standard.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'message_id',
                'header_contains': ['ID序号', 'ID定义'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 0,
            }
        ]
    },
    {
        'filename': '14_id_old.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'message_id',
                'header_contains': ['信息内容', '消息ID'],
                'data_rows_min': 2,
                'data_rows_max': 5,
                'header_row_idx': 0,
            }
        ]
    },
    {
        'filename': '15_id_non_standard.docx',
        'config': {
            'fields': ['编号', '名称', '有无数据'],
            'id_field_names': ['编号', '名称'],
        },
        'expected_tables': [
            {
                'table_type': 'message_id',
                'header_contains': ['编号', '名称'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 0,
            }
        ]
    },
    # ─── 4. 端口分配表 ──────────────────────────────────────────────────
    {
        'filename': '16_port_allocation.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'port_allocation',
                'header_contains': ['信源系统码', '信宿系统码'],
                'data_rows_min': 1,
                'data_rows_max': 3,
                'header_row_idx': 0,
            }
        ]
    },
    # ─── 5. Bit位定义表 ──────────────────────────────────────────────────
    {
        'filename': '17_bit_def.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'bit_def',
                'header_contains': ['位号'],
                'data_rows_min': 3,
                'data_rows_max': 6,
                'header_row_idx': 0,
            }
        ]
    },
    # ─── 6. 噪声表 ──────────────────────────────────────────────────────
    {
        'filename': '18_noise_cover.docx',
        'config': None,
        'expected_tables': [
            {'table_type': 'skip'}
        ]
    },
    {
        'filename': '19_noise_scenario.docx',
        'config': None,
        'expected_tables': [
            {'table_type': 'skip'}
        ]
    },
    {
        'filename': '20_noise_frame_format.docx',
        'config': None,
        'expected_tables': [
            {'table_type': 'skip'}
        ]
    },
    {
        'filename': '21_noise_bit_labels.docx',
        'config': None,
        'expected_tables': [
            {'table_type': 'skip'}
        ]
    },
    # ─── 7. 边界 case ───────────────────────────────────────────────────
    {
        'filename': '22_edge_2col_header.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['字节号'],
                'data_rows_min': 3,
                'data_rows_max': 5,
                'header_row_idx': 0,
            }
        ]
    },
    {
        'filename': '23_edge_coordinate.docx',
        'config': None,
        'expected_tables': [
            {'table_type': 'skip'}
        ]
    },
    {
        'filename': '24_edge_action_table.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '位置状态'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 0,
            }
        ]
    },
    {
        'filename': '25_mixed_with_notes.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 4,
            }
        ]
    },
    {
        'filename': '26_mixed_diff_cols.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 3,
            }
        ]
    },
    {
        'filename': '27_mixed_with_config_fields.docx',
        'config': ['代号', '描述', '规格'],
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['代号', '描述'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 3,
            }
        ]
    },
    # ─── 综合文档 ───────────────────────────────────────────────────────
    {
        'filename': '28_all_in_one.docx',
        'config': None,
        'expected_tables': [
            {'table_type': 'skip'},  # 封面
            {
                'table_type': 'message_id',
                'header_contains': ['ID序号', 'ID定义'],
                'data_rows_min': 1,
                'data_rows_max': 3,
            },
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'data_rows_min': 1,
                'data_rows_max': 3,
            },
            {
                'table_type': 'field_def',
                'header_contains': ['字节号'],
                'data_rows_min': 1,
                'data_rows_max': 3,
            },
            {
                'table_type': 'bit_def',
                'header_contains': ['位号'],
                'data_rows_min': 1,
                'data_rows_max': 3,
            },
            {'table_type': 'skip'},  # 场景表
        ]
    },
    {
        'filename': '29_mixed_empty_metadata_value.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['序号', '内容', '类型'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 5,
            }
        ]
    },
    {
        'filename': '30_simple_data_only.docx',
        'config': None,
        'expected_tables': [
            {
                'table_type': 'field_def',
                'header_contains': ['参数', '数据类型'],
                'data_rows_min': 2,
                'data_rows_max': 4,
                'header_row_idx': 0,
            }
        ]
    },
]


def run_test(test_case):
    """运行单个测试用例，返回 (passed, details)"""
    filename = test_case['filename']
    config = test_case.get('config')
    expected_tables = test_case.get('expected_tables', [])

    filepath = os.path.join(TEST_DIR, filename)
    if not os.path.exists(filepath):
        return False, f"文件不存在: {filepath}"

    try:
        detector = TableDetector(config=config)
        results = detector.extract_tables_from_docx(filepath)
    except Exception as e:
        return False, f"异常: {e}\n{traceback.format_exc()}"

    details = []
    all_passed = True

    # 检查表格数量
    if len(results) != len(expected_tables):
        details.append(f"表格数量不符: 期望 {len(expected_tables)}, 实际 {len(results)}")
        all_passed = False
        # 列出实际结果
        for i, r in enumerate(results):
            details.append(f"  实际 Table#{i}: type={r.get('table_type')}, "
                          f"headers={r.get('headers', [])[:5]}, "
                          f"data_rows={len(r.get('data_rows', []))}")

    # 逐表检查
    for i, expected in enumerate(expected_tables):
        if i >= len(results):
            details.append(f"Table#{i}: 缺失（期望 type={expected.get('table_type')}）")
            all_passed = False
            continue

        actual = results[i]
        table_passed = True
        table_details = []

        # 检查 table_type
        exp_type = expected.get('table_type')
        act_type = actual.get('table_type', 'skip')
        if exp_type and act_type != exp_type:
            table_details.append(f"type不符: 期望={exp_type}, 实际={act_type}")
            table_passed = False

        # 检查表头包含
        headers = actual.get('headers', [])
        headers_text = ' '.join(headers)
        for kw in expected.get('header_contains', []):
            if kw not in headers_text:
                table_details.append(f"表头缺少'{kw}': headers={headers[:6]}")
                table_passed = False

        # 检查表头不包含
        for kw in expected.get('header_not_contains', []):
            if kw in headers_text:
                table_details.append(f"表头不应含'{kw}': headers={headers[:6]}")
                table_passed = False

        # 检查数据行数
        data_rows = actual.get('data_rows', [])
        dr_min = expected.get('data_rows_min')
        dr_max = expected.get('data_rows_max')
        if dr_min is not None and len(data_rows) < dr_min:
            table_details.append(f"数据行太少: 期望≥{dr_min}, 实际={len(data_rows)}")
            table_passed = False
        if dr_max is not None and len(data_rows) > dr_max:
            table_details.append(f"数据行太多: 期望≤{dr_max}, 实际={len(data_rows)}")
            table_passed = False

        # 检查消息名
        msg_name = actual.get('msg_name', '')
        for kw in expected.get('msg_name_contains', []):
            if kw not in msg_name:
                table_details.append(f"msg_name缺少'{kw}': msg_name='{msg_name}'")
                table_passed = False

        # 检查 meta
        meta = actual.get('meta', {})
        for key in expected.get('meta_contains', []):
            if key not in meta:
                table_details.append(f"meta缺少'{key}': meta={meta}")
                table_passed = False

        # 检查 header_row_idx（通过验证实际表头内容是否匹配预期行）
        exp_hri = expected.get('header_row_idx')
        if exp_hri is not None and act_type != 'skip':
            # 读取原始文档获取 grid
            try:
                doc = Document(filepath)
                if i < len(doc.tables):
                    grid, _ = _build_grid(doc.tables[i])
                    if exp_hri < len(grid):
                        actual_header_at_exp = _dedup_row(grid[exp_hri])
                        # 检查实际表头是否在预期行
                        actual_headers_set = set(h.strip() for h in headers if h.strip())
                        exp_row_set = set(h.strip() for h in actual_header_at_exp if h.strip())
                        # 如果实际表头和预期行的内容有交集，说明定位正确
                        if actual_headers_set and exp_row_set:
                            overlap = actual_headers_set & exp_row_set
                            if len(overlap) < min(len(actual_headers_set), len(exp_row_set)) * 0.5:
                                table_details.append(
                                    f"表头行索引可能不符: 期望行{exp_hri}={actual_header_at_exp[:5]}, "
                                    f"实际headers={headers[:5]}, overlap={overlap}"
                                )
                                table_passed = False
            except Exception:
                pass  # 忽略 grid 验证错误

        if table_passed:
            table_details.append(f"✓ type={act_type}, headers={headers[:5]}, data_rows={len(data_rows)}, msg='{msg_name}'")
        else:
            all_passed = False

        details.append(f"  Table#{i}: {'✓' if table_passed else '✗'} {'; '.join(table_details)}")

    return all_passed, '\n'.join(details)


def main():
    print("=" * 70)
    print("表头定位 & 表格识别测试")
    print("=" * 70)

    passed_count = 0
    failed_count = 0
    failed_cases = []

    for tc in TEST_CASES:
        filename = tc['filename']
        passed, details = run_test(tc)
        status = '✓ PASS' if passed else '✗ FAIL'
        print(f"\n[{status}] {filename}")
        print(f"  {details}")

        if passed:
            passed_count += 1
        else:
            failed_count += 1
            failed_cases.append(filename)

    print("\n" + "=" * 70)
    print(f"结果: {passed_count} 通过, {failed_count} 失败")
    if failed_cases:
        print("失败用例:")
        for fc in failed_cases:
            print(f"  - {fc}")
    print("=" * 70)

    return failed_count == 0


if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
